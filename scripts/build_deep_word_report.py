import json
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
FIG_DIR = REPORT_DIR / "figures"
TABLE_DIR = REPORT_DIR / "tables"
OUT_DIR = REPORT_DIR / "word"
OUT_PATH = OUT_DIR / "基于ResNet交通标志识别的对抗扰动攻击与输入防御深度实验报告.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 90, 90)
BLACK = RGBColor(0, 0, 0)
ACCENT = RGBColor(31, 58, 95)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
TABLE_WIDTH_DXA = 9360


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def num(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFC7D1", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths_dxa, font_size=9.2, header_fill=LIGHT_GRAY):
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    if table.rows:
        set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.10
                for run in p.runs:
                    set_run_font(run, size=font_size)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths_dxa, caption=None, font_size=9.2, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = str(text)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths_dxa, font_size=font_size, header_fill=header_fill)
    if caption:
        add_caption(doc, caption)
    else:
        doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9.2, color=MUTED)


def add_para(doc, text, after=6, before=0, bold=False, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, name="Cambria Math", east_asia="Microsoft YaHei", size=10.5, color=NAVY)
    return p


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9.2, color=RGBColor(40, 40, 40))
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=ACCENT)
    p.paragraph_format.space_after = Pt(3)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=BLACK)
    style_table(table, [TABLE_WIDTH_DXA], font_size=10.2, header_fill=CALLOUT)
    set_cell_shading(cell, CALLOUT)
    doc.add_paragraph()


def add_figure(doc, filename, caption, width=6.05, note=None):
    path = FIG_DIR / filename
    if not path.exists():
        add_para(doc, f"[缺失图像：{filename}]", color=RGBColor(155, 28, 28))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", filename)
    add_caption(doc, caption)
    if note:
        p_note = doc.add_paragraph()
        p_note.paragraph_format.space_after = Pt(6)
        r = p_note.add_run(note)
        set_run_font(r, size=9.5, italic=True, color=MUTED)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("ResNet交通标志识别对抗鲁棒性实验报告 | 第 ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(footer)
    r = footer.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def diagram_boxes(path, title, boxes, arrows):
    fig, ax = plt.subplots(figsize=(11, 4.8))
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.text(0.02, 0.94, title, fontsize=15, fontweight="bold", color="#0B2545")
    for key, text, x, y, w, h, fill in boxes:
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.018,rounding_size=0.025",
            linewidth=1.2,
            edgecolor="#355070",
            facecolor=fill,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=10, color="#111111")
    loc = {key: (x, y, w, h) for key, _, x, y, w, h, _ in boxes}
    for start, end in arrows:
        sx, sy, sw, sh = loc[start]
        ex, ey, ew, eh = loc[end]
        arrow = FancyArrowPatch(
            (sx + sw, sy + sh / 2),
            (ex, ey + eh / 2),
            arrowstyle="-|>",
            mutation_scale=14,
            linewidth=1.2,
            color="#4A5568",
        )
        ax.add_patch(arrow)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def generate_diagrams():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    diagram_boxes(
        FIG_DIR / "fig_23_experiment_pipeline.png",
        "End-to-end experiment chain",
        [
            ("data", "GTSRB\nraw images", 0.03, 0.56, 0.13, 0.18, "#E8EEF5"),
            ("prep", "Resize +\nAugment + Norm", 0.20, 0.56, 0.14, 0.18, "#F2F4F7"),
            ("resnet", "ResNet18\nclassifier", 0.38, 0.56, 0.14, 0.18, "#E8EEF5"),
            ("clean", "Clean\nmetrics", 0.56, 0.71, 0.13, 0.15, "#F7FAFC"),
            ("attack", "FGSM / PGD\nLinf attack", 0.56, 0.45, 0.13, 0.15, "#FFF4E6"),
            ("percept", "Perturbation\nPSNR + Linf", 0.73, 0.45, 0.13, 0.15, "#FFF7ED"),
            ("defense", "Input defense\nBlur / Median / JPEG", 0.73, 0.71, 0.18, 0.15, "#ECFDF5"),
        ],
        [("data", "prep"), ("prep", "resnet"), ("resnet", "clean"), ("resnet", "attack"), ("attack", "percept"), ("percept", "defense")],
    )
    diagram_boxes(
        FIG_DIR / "fig_24_resnet_recognition_flow.png",
        "ResNet18 recognition mechanism",
        [
            ("x", "Image tensor\n3 x 64 x 64", 0.03, 0.56, 0.13, 0.18, "#E8EEF5"),
            ("stem", "3x3 Conv\nBN + ReLU", 0.20, 0.56, 0.13, 0.18, "#F2F4F7"),
            ("blocks", "Residual stages\n[2,2,2,2]", 0.37, 0.56, 0.14, 0.18, "#E8EEF5"),
            ("gap", "Global avg\npooling", 0.55, 0.56, 0.12, 0.18, "#F2F4F7"),
            ("logits", "FC layer\n43 logits", 0.71, 0.56, 0.12, 0.18, "#E8EEF5"),
            ("softmax", "Softmax\nclass prob.", 0.87, 0.56, 0.10, 0.18, "#F7FAFC"),
            ("res", "Residual block:\ny = F(x, W) + x", 0.35, 0.18, 0.31, 0.17, "#FFF7ED"),
        ],
        [("x", "stem"), ("stem", "blocks"), ("blocks", "gap"), ("gap", "logits"), ("logits", "softmax")],
    )
    diagram_boxes(
        FIG_DIR / "fig_25_attack_threat_model.png",
        "White-box untargeted perturbation attack",
        [
            ("clean", "Clean input x\ncorrect class y", 0.05, 0.56, 0.16, 0.18, "#E8EEF5"),
            ("grad", "Backprop to input\nsign(grad_x L)", 0.28, 0.56, 0.17, 0.18, "#FFF4E6"),
            ("proj", "Project to\nLinf epsilon ball", 0.52, 0.56, 0.18, 0.18, "#FFF7ED"),
            ("adv", "Adversarial x_adv\nwrong prediction", 0.78, 0.56, 0.17, 0.18, "#FEE2E2"),
            ("note", "Actual image is x_adv.\nDelta heatmaps are amplified analysis views.", 0.28, 0.20, 0.42, 0.17, "#F4F6F9"),
        ],
        [("clean", "grad"), ("grad", "proj"), ("proj", "adv")],
    )
    diagram_boxes(
        FIG_DIR / "fig_26_defense_data_link.png",
        "Input preprocessing defense data link",
        [
            ("adv", "Adversarial\nimage", 0.07, 0.56, 0.15, 0.18, "#FEE2E2"),
            ("pixel", "Denormalize to\npixel space", 0.30, 0.56, 0.15, 0.18, "#F2F4F7"),
            ("pre", "Preprocess\nBlur / Median / JPEG", 0.53, 0.56, 0.19, 0.18, "#ECFDF5"),
            ("norm", "Normalize again\nsame mean/std", 0.79, 0.56, 0.15, 0.18, "#F2F4F7"),
            ("model", "Frozen ResNet18\nprediction", 0.36, 0.20, 0.20, 0.17, "#E8EEF5"),
        ],
        [("adv", "pixel"), ("pixel", "pre"), ("pre", "norm"), ("norm", "model")],
    )


def add_cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("模式识别与智能系统\n期末大作业深度实验报告")
    set_run_font(r, size=22, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("基于 ResNet 交通标志识别的对抗扰动攻击与输入防御鲁棒性研究")
    set_run_font(r, size=17, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("从深度学习识别原理到 FGSM/PGD 攻击、扰动可感知性分析与防御消融")
    set_run_font(r, size=12, color=MUTED)

    rows = [
        ("课程名称", "模式识别与智能系统"),
        ("项目方向", "深度学习识别、对抗扰动攻击、模型鲁棒防御"),
        ("核心数据集", "GTSRB German Traffic Sign Recognition Benchmark，43 类交通标志"),
        ("模型主线", "ResNet18 小尺寸图像适配版，输出 43 类 logits"),
        ("实验主线", "Clean recognition -> FGSM/PGD attack -> perturbation analysis -> input defense"),
        ("完成状态", "基础识别、攻击、防御和扰动可感知性分析已完成；对抗训练为后续验证项"),
        ("小组成员", "待填写"),
        ("提交日期", "2026 年 6 月"),
    ]
    add_table(doc, ["项目", "内容"], rows, [1900, 7460], caption=None)
    doc.add_page_break()


def add_abstract_and_summary(doc, data):
    doc.add_heading("摘要", level=1)
    baseline = data["baseline"]
    attack_df = data["attack"]
    defense_df = data["defense"]
    pert_df = data["perturbation"]

    pgd003 = attack_df[(attack_df.attack == "pgd") & (attack_df.epsilon.round(3) == 0.03)].iloc[0]
    jpeg_pgd003 = defense_df[
        (defense_df.attack == "pgd") & (defense_df.epsilon.round(3) == 0.03) & (defense_df.defense == "jpeg_compression")
    ].iloc[0]
    pert_pgd003 = pert_df[
        (pert_df.attack == "pgd") & (pert_df.epsilon_normalized_space.round(3) == 0.03)
    ].iloc[0]

    add_para(
        doc,
        "本报告围绕交通标志识别模型的鲁棒性展开，不只完成普通图像分类复现，而是构建“深度学习识别原理 - ResNet18 基础模型 - 对抗扰动攻击 - 输入防御 - 消融对比”的完整实验链路。项目使用 GTSRB 数据集，先训练小尺寸图像适配版 ResNet18，随后在白盒、无目标、Linf 约束条件下实现 FGSM 与 PGD 攻击，并使用 Gaussian Blur、Median Filter、JPEG Compression 三类输入预处理方法进行防御评估。",
    )
    add_para(
        doc,
        f"当前已完成结果显示：ResNet18 在正常测试集上的准确率为 {pct(baseline['accuracy'])}；在 PGD epsilon=0.03 攻击下，对抗样本准确率下降至 {pct(pgd003.adversarial_accuracy)}，攻击成功率为 {pct(pgd003.attack_success_rate)}。扰动并不是肉眼明显的大面积改图，epsilon=0.03 在反归一化像素空间的平均 Linf 约为 {num(pert_pgd003.mean_linf_pixel_space, 4)}，约等于 2/255，平均 PSNR 为 {num(pert_pgd003.mean_psnr_db, 2)} dB。输入防御中 JPEG Compression 表现最好，在 PGD epsilon=0.03 下将准确率恢复到 {pct(jpeg_pgd003.defended_accuracy)}。",
    )
    p = doc.add_paragraph()
    r = p.add_run("关键词：")
    set_run_font(r, size=11, bold=True)
    r = p.add_run("交通标志识别；ResNet18；GTSRB；对抗样本；FGSM；PGD；Linf 扰动；PSNR；输入预处理防御；鲁棒性")
    set_run_font(r, size=11)

    doc.add_heading("核心结论速览", level=1)
    rows = [
        ("基础识别", "ResNet18 clean test accuracy", pct(baseline["accuracy"]), "说明基础分类能力成立"),
        ("攻击效果", "PGD eps=0.03 adversarial accuracy", pct(pgd003.adversarial_accuracy), "小扰动即可显著破坏识别"),
        ("真实扰动", "PGD eps=0.03 pixel-space Linf / PSNR", f"{num(pert_pgd003.mean_linf_pixel_space, 4)} / {num(pert_pgd003.mean_psnr_db, 2)} dB", "扰动是小幅像素扰动，热力图为放大视图"),
        ("防御效果", "JPEG defense under PGD eps=0.03", pct(jpeg_pgd003.defended_accuracy), "输入预处理可部分恢复准确率"),
    ]
    add_table(doc, ["环节", "关键指标", "结果", "解释"], rows, [1500, 2800, 2100, 2960], "表 1  核心实验结论汇总")
    add_callout(
        doc,
        "重要澄清：为什么之前的攻击图看起来“太明显”？",
        "旧图中的 Noise/Delta 列是为了让扰动可见而做的归一化热力图或放大图，不是实际送入模型的对抗图像。实际对抗样本是 Original 与 Adversarial 两列的差异，epsilon=0.03 在像素空间平均约 0.008，即大约 2/255。新报告增加 PSNR、Linf 和 Delta x30 标注，明确区分真实图像与放大分析视图。",
    )
    doc.add_page_break()


def add_section_1(doc):
    doc.add_heading("1. 任务背景与实验总链路", level=1)
    doc.add_heading("1.1 任务背景", level=2)
    add_para(
        doc,
        "交通标志识别是智能交通和自动驾驶视觉感知中的基础任务。系统需要从道路图像中识别限速、禁止通行、让行、转向等标志，并将识别结果交给后续规划或告警模块。深度卷积网络可以在标准测试集上取得很高准确率，但其安全性不能只由 clean accuracy 判断；在微小扰动、压缩失真、光照变化或恶意攻击下，模型可能输出错误类别。",
    )
    add_para(
        doc,
        "因此，本项目将普通分类任务扩展为鲁棒性实验：先建立可信的识别基线，再构造受限扰动攻击，最后比较防御策略。这样的设计比单纯复现分类模型更能体现技术深度，也更适合展示模型在真实安全场景中的风险和改进空间。",
    )

    doc.add_heading("1.2 研究问题", level=2)
    for item in [
        "ResNet18 是否能够在 GTSRB 交通标志任务上建立稳定的 clean recognition 基线？",
        "在白盒条件下，FGSM 与 PGD 这类梯度攻击能否用小幅 Linf 扰动显著降低识别准确率？",
        "epsilon 增大时，攻击成功率、对抗准确率、置信度下降和可感知性指标如何变化？",
        "不重新训练模型的输入预处理防御能否恢复部分对抗样本准确率？哪一种防御更有效？",
        "哪些实验已经有量化证据，哪些仍应作为后续补充，而不能在报告中虚构结果？",
    ]:
        add_numbered(doc, item)

    doc.add_heading("1.3 总体实验链路", level=2)
    add_figure(doc, "fig_23_experiment_pipeline.png", "图 1  项目总实验链路：识别、攻击、扰动分析、防御对比", width=6.25)
    rows = [
        ("E00", "数据检查", "类别分布、图像尺寸、样例图", "dataset_summary.csv、样例网格图"),
        ("E01", "基础识别", "训练 ResNet18 并评估 clean accuracy", "训练曲线、混淆矩阵、错误样例"),
        ("E02", "扰动攻击", "FGSM/PGD + epsilon 扫描", "攻击曲线、对抗样例、攻击指标"),
        ("E02b", "可感知性分析", "像素空间 Linf、MSE、PSNR", "PSNR 柱状图、Delta x30 对比图"),
        ("E03", "输入防御", "Blur/Median/JPEG 防御消融", "防御曲线、恢复样例、防御指标"),
        ("E04", "对抗训练", "代码已准备，未完成验证", "不纳入正式结论"),
    ]
    add_table(doc, ["编号", "实验环节", "实验目的", "保存产物"], rows, [1100, 1700, 3300, 3260], "表 2  实验链路与产物设计")


def add_section_2(doc, data):
    doc.add_heading("2. 深度学习识别原理与 ResNet 模型结构", level=1)
    doc.add_heading("2.1 监督分类的数学形式", level=2)
    add_para(
        doc,
        "交通标志识别可形式化为 43 类监督分类问题。输入图像经过预处理后表示为张量 x in R^(3 x 64 x 64)，标签 y in {0,...,42}。神经网络 f_theta 将输入映射为 43 维 logits z，softmax 将 logits 转换为类别概率。训练目标是最小化交叉熵损失，使真实类别的预测概率尽可能高。",
    )
    add_formula(doc, "z = f_theta(x),    p_i = exp(z_i) / sum_j exp(z_j),    L(x, y; theta) = -log p_y")
    add_para(
        doc,
        "在反向传播中，损失函数对模型参数 theta 求梯度，优化器根据梯度更新卷积核、BatchNorm 参数和全连接层参数。攻击实验中同样利用反向传播，但梯度对象从模型参数转为输入图像 x，这正是白盒对抗攻击能够构造扰动的原因。",
    )

    doc.add_heading("2.2 卷积网络识别机制", level=2)
    rows = [
        ("卷积层", "通过共享卷积核提取局部边缘、颜色、纹理和形状模式", "保留空间结构，参数量低于全连接"),
        ("BatchNorm", "对中间特征做标准化并学习缩放平移", "稳定训练，加快收敛"),
        ("ReLU", "非线性激活 max(0, x)", "提高模型表达能力"),
        ("池化/下采样", "逐层扩大感受野", "从局部纹理过渡到整体标志形状"),
        ("全局平均池化", "对最后特征图做空间汇聚", "降低参数量，输出类别特征"),
        ("Softmax", "将 logits 转换为概率分布", "用于分类决策与置信度分析"),
    ]
    add_table(doc, ["模块", "作用", "在交通标志识别中的意义"], rows, [1800, 3600, 3960], "表 3  CNN 识别模块及作用")
    add_figure(doc, "fig_24_resnet_recognition_flow.png", "图 2  ResNet18 交通标志识别数据流与残差块结构", width=6.25)

    doc.add_heading("2.3 ResNet 残差学习原理", level=2)
    add_para(
        doc,
        "普通深层网络在层数增加后可能出现退化问题：理论上更深的网络表达能力更强，但实际优化更难，训练误差可能反而升高。ResNet 的核心是残差连接，让网络学习残差函数 F(x, W)，而不是直接学习完整映射 H(x)。残差块输出为：",
    )
    add_formula(doc, "y = F(x, W) + x")
    add_para(
        doc,
        "该结构提供了一条近似恒等映射的梯度通路，使低层特征和梯度可以更顺畅地穿过深层网络。对于交通标志这种小目标分类任务，早期边缘和颜色信息很重要，因此本项目对 torchvision ResNet18 做了小图像适配：第一层使用 3x3 stride=1 卷积，去除原始 ImageNet 版本中的 maxpool，避免 64x64 输入过早丢失空间细节。",
    )

    rows = [
        ("Input", "3 x 64 x 64", "归一化后的 RGB 图像"),
        ("Stem", "Conv 3x3, 64, stride=1 + BN + ReLU", "替代原 ResNet 的 7x7 stride=2"),
        ("MaxPool", "Identity", "小图像不做初始最大池化"),
        ("Layer1", "2 residual blocks, 64 channels", "低级纹理与边缘"),
        ("Layer2", "2 residual blocks, 128 channels", "局部形状组合"),
        ("Layer3", "2 residual blocks, 256 channels", "类别相关形状特征"),
        ("Layer4", "2 residual blocks, 512 channels", "高层语义特征"),
        ("Classifier", "GlobalAvgPool + Linear(512, 43)", "输出 43 类 logits"),
    ]
    add_table(doc, ["阶段", "结构", "说明"], rows, [1700, 3300, 4360], "表 4  本项目 ResNet18 模型结构")

    doc.add_heading("2.4 识别数据链路", level=2)
    rows = [
        ("原始图片", "GTSRB RGB 图像，尺寸不一", "data/raw"),
        ("Resize", "统一缩放到 64 x 64", "保证 batch 输入尺寸一致"),
        ("训练增强", "RandomRotation、ColorJitter、RandomAffine", "提升光照、角度和小幅形变泛化"),
        ("ToTensor", "像素从 [0,255] 转为 [0,1]", "后续归一化和模型输入"),
        ("Normalize", "mean=(0.3337,0.3064,0.3171)，std=(0.2672,0.2564,0.2629)", "攻击 epsilon 在该归一化空间定义"),
        ("Forward", "ResNet18 输出 logits", "得到 43 类未归一化分数"),
        ("Loss/Metric", "CrossEntropy、Accuracy、Macro F1", "训练与评估依据"),
    ]
    add_table(doc, ["链路阶段", "处理内容", "用途"], rows, [1900, 4200, 3260], "表 5  基础识别数据链路")


def add_section_3(doc, data):
    doc.add_heading("3. 数据集、基础识别实验与结果分析", level=1)
    overview = data["dataset"]
    baseline = data["baseline"]
    train_log = data["train_log"]
    per_class = data["per_class"]

    doc.add_heading("3.1 数据集统计与展示", level=2)
    rows = [
        ("类别数", 43),
        ("训练集图片数", overview["train_images"]),
        ("测试集图片数", overview["test_images"]),
        ("总图片数", overview["total_images"]),
        ("最小宽度 / 最大宽度", f"{overview['min_width']} / {overview['max_width']}"),
        ("最小高度 / 最大高度", f"{overview['min_height']} / {overview['max_height']}"),
    ]
    add_table(doc, ["项目", "数值"], rows, [2600, 6760], "表 6  GTSRB 数据集基本统计")
    add_figure(doc, "fig_01_class_distribution.png", "图 3  GTSRB 类别分布统计", width=6.05)
    add_figure(doc, "fig_02_image_size_distribution.png", "图 4  GTSRB 图像尺寸分布统计", width=5.95)
    add_figure(doc, "fig_03_train_samples.png", "图 5  训练集交通标志样例", width=6.20)
    add_figure(doc, "fig_04_test_samples.png", "图 6  测试集交通标志样例", width=6.20)

    doc.add_heading("3.2 基础识别实验目的与配置", level=2)
    add_para(
        doc,
        "基础识别实验的目标是建立一个性能足够稳定的 clean model，作为后续攻击与防御实验的被攻击对象。如果基础模型本身识别能力不足，则攻击成功不能说明鲁棒性问题。因此先保证模型在正常测试集上达到较高准确率，再进行对抗实验。",
    )
    rows = [
        ("模型", "ResNet18，小图像适配版"),
        ("输入尺寸", "64 x 64"),
        ("类别数", "43"),
        ("训练/验证划分", "训练集内部 85% / 15%"),
        ("Batch size", "256"),
        ("Epoch", "30"),
        ("优化器", "AdamW"),
        ("学习率", "0.001"),
        ("权重衰减", "0.0001"),
        ("调度器", "Cosine learning rate schedule"),
        ("损失函数", "CrossEntropyLoss"),
        ("保存策略", "保存 best_model.pth、last_model.pth、训练曲线、混淆矩阵、预测样例"),
    ]
    add_table(doc, ["配置项", "设置"], rows, [2600, 6760], "表 7  ResNet18 基础识别实验配置")

    doc.add_heading("3.3 训练过程与 clean 测试结果", level=2)
    best_epoch_row = train_log.sort_values("val_acc", ascending=False).iloc[0]
    last_epoch_row = train_log.sort_values("epoch").iloc[-1]
    rows = [
        ("测试集 Accuracy", pct(baseline["accuracy"])),
        ("Precision macro", num(baseline["precision_macro"])),
        ("Recall macro", num(baseline["recall_macro"])),
        ("F1 macro", num(baseline["f1_macro"])),
        ("Best epoch", int(baseline["best_epoch"])),
        ("Best validation accuracy", pct(baseline["best_val_accuracy"])),
        ("参数量", f"{baseline['parameter_count']:,}"),
        ("训练耗时", f"{baseline['training_seconds'] / 60:.1f} min"),
        ("训练日志最高 val_acc epoch", int(best_epoch_row.epoch)),
        ("最后 epoch train_acc / val_acc", f"{pct(last_epoch_row.train_acc)} / {pct(last_epoch_row.val_acc)}"),
    ]
    add_table(doc, ["指标", "数值"], rows, [3300, 6060], "表 8  ResNet18 clean 测试与训练摘要")
    add_figure(doc, "fig_05_baseline_loss_curve.png", "图 7  ResNet18 训练损失曲线", width=5.70)
    add_figure(doc, "fig_06_baseline_accuracy_curve.png", "图 8  ResNet18 训练/验证准确率曲线", width=5.70)
    add_figure(doc, "fig_07_baseline_confusion_matrix.png", "图 9  ResNet18 测试集混淆矩阵", width=5.90)
    add_figure(doc, "fig_08_baseline_per_class_accuracy.png", "图 10  ResNet18 各类别测试准确率", width=6.05)

    doc.add_heading("3.4 类别级结果与错误样例分析", level=2)
    low = per_class.sort_values("accuracy", ascending=True).head(6)
    high = per_class.sort_values(["accuracy", "total"], ascending=[False, False]).head(6)
    rows = []
    for _, row in low.iterrows():
        rows.append((int(row.class_id), row.class_name, int(row.total), int(row.correct), pct(row.accuracy)))
    add_table(doc, ["Class", "类别名称", "样本数", "正确数", "准确率"], rows, [900, 3300, 1500, 1500, 2160], "表 9  准确率较低类别示例")
    rows = []
    for _, row in high.iterrows():
        rows.append((int(row.class_id), row.class_name, int(row.total), int(row.correct), pct(row.accuracy)))
    add_table(doc, ["Class", "类别名称", "样本数", "正确数", "准确率"], rows, [900, 3300, 1500, 1500, 2160], "表 10  准确率较高类别示例")
    add_figure(doc, "fig_09_baseline_correct_samples.png", "图 11  基础模型正确识别样例", width=6.20)
    add_figure(doc, "fig_10_baseline_wrong_samples.png", "图 12  基础模型错误识别样例", width=6.20)
    add_para(
        doc,
        "从类别级指标和错误样例可以看出，clean 准确率很高，但模型仍会在低光照、模糊、遮挡、类别形状相近或图像质量较差时出错。这为后续攻击实验提供了直观背景：对抗扰动并不是凭空制造脆弱性，而是利用模型决策边界附近的敏感方向放大错误风险。",
    )

    doc.add_heading("3.5 基础识别部分的对比与消融说明", level=2)
    rows = [
        ("已完成对照", "clean test vs adversarial test", "已量化", "后续攻击章节给出同一模型在 clean/adv 条件下的准确率差异"),
        ("已完成分析", "训练曲线 vs 测试结果", "已量化", "验证模型不是随机失败，clean 基线成立"),
        ("未完成消融", "原始 ResNet18 stem vs 小图像适配 stem", "未训练", "不能伪造指标；如后续补跑，可比较 7x7+maxpool 与 3x3+no maxpool"),
        ("未完成消融", "无数据增强 vs 当前增强", "未训练", "可作为答辩时的后续实验计划"),
    ]
    add_table(doc, ["类别", "对比/消融项", "当前状态", "报告处理方式"], rows, [1400, 2700, 1600, 3660], "表 11  基础识别部分对比与消融完整性说明")


def attack_rows(df):
    rows = []
    for _, row in df.sort_values(["attack", "epsilon"]).iterrows():
        rows.append(
            [
                row.attack.upper(),
                num(row.epsilon, 3),
                pct(row.clean_accuracy),
                pct(row.adversarial_accuracy),
                pct(row.attack_success_rate),
                num(row.mean_confidence_drop, 4),
            ]
        )
    return rows


def add_section_4(doc, data):
    doc.add_heading("4. 对抗扰动攻击原理、实验与可感知性分析", level=1)
    attack_df = data["attack"]
    pert_df = data["perturbation"]

    doc.add_heading("4.1 威胁模型与攻击目标", level=2)
    add_para(
        doc,
        "本项目攻击设置为白盒、无目标、Linf 约束攻击。白盒表示攻击者知道模型结构和参数，可以计算损失对输入的梯度；无目标表示攻击只要求模型从正确类别变为任意错误类别；Linf 约束表示每个像素通道的最大扰动幅度被 epsilon 限制。",
    )
    add_formula(doc, "x_adv = arg max_{x' in B_inf(x, epsilon)} L(f_theta(x'), y)")
    rows = [
        ("攻击知识", "白盒", "可访问 ResNet18 权重与输入梯度"),
        ("攻击目标", "无目标攻击", "只要求预测类别不等于真实标签"),
        ("扰动约束", "Linf ball", "每个像素通道扰动幅度受 epsilon 限制"),
        ("样本范围", "分层抽样 3000 张测试图", "攻击与防御使用同一索引，保证可比"),
        ("模型状态", "冻结 best_model.pth", "攻击和防御均不更新模型参数"),
    ]
    add_table(doc, ["维度", "设置", "解释"], rows, [1800, 2600, 4960], "表 12  对抗攻击威胁模型")
    add_figure(doc, "fig_25_attack_threat_model.png", "图 13  白盒无目标扰动攻击示意图", width=6.25)

    doc.add_heading("4.2 FGSM 与 PGD 数学结构", level=2)
    add_para(
        doc,
        "FGSM 是单步攻击。它假设在输入邻域内，损失函数可用一阶线性近似描述，因此沿输入梯度符号方向走一步即可快速增大损失。其优点是计算快，缺点是搜索不充分。",
    )
    add_formula(doc, "FGSM:    x_adv = clip_{[data_min,data_max]}(x + epsilon * sign(grad_x L(f_theta(x), y)))")
    add_para(
        doc,
        "PGD 是迭代攻击。它从 epsilon 邻域内随机初始化，每一步沿梯度符号方向更新，再投影回 Linf epsilon ball。PGD 相当于在受限邻域内多次寻找更强的损失上升方向，通常比 FGSM 更接近强攻击基线。",
    )
    add_formula(doc, "PGD:    x_{t+1} = Proj_{B_inf(x,epsilon)}(x_t + alpha * sign(grad_x L(f_theta(x_t), y)))")

    rows = [
        ("FGSM", "单步", "epsilon", "0.005, 0.01, 0.02, 0.03, 0.05", "速度快，攻击强度随 epsilon 增大"),
        ("PGD", "7 步迭代", "epsilon + alpha", "epsilon 同上，alpha=0.005", "更强但更耗时"),
    ]
    add_table(doc, ["方法", "更新次数", "关键超参", "本实验设置", "特点"], rows, [1100, 1400, 1600, 3000, 2260], "表 13  FGSM 与 PGD 方法对比")

    doc.add_heading("4.3 攻击实验数据链路与配置", level=2)
    rows = [
        ("输入", "测试集图像 -> Resize -> ToTensor -> Normalize", "与基础识别保持一致"),
        ("前向", "ResNet18 输出 clean prediction", "记录 clean correctness 与 confidence"),
        ("反向", "对输入求 grad_x L", "只用于生成扰动，不更新模型"),
        ("投影", "限制到 data_min/data_max 与 epsilon ball", "保证图像仍处于合法归一化范围"),
        ("评估", "用同一 ResNet18 预测 x_adv", "记录 adversarial accuracy、success rate、confidence drop"),
        ("保存", "CSV + 曲线 + 对抗样例 + 像素空间可感知性表", "报告可复现、可展示"),
    ]
    add_table(doc, ["阶段", "数据处理", "目的"], rows, [1700, 4600, 3060], "表 14  对抗攻击数据链路")
    rows = [
        ("模型 checkpoint", "results/01_baseline/resnet18/checkpoints/best_model.pth"),
        ("评估样本", "3000 张分层抽样测试图"),
        ("Batch size", "128"),
        ("FGSM epsilon", "0.005, 0.01, 0.02, 0.03, 0.05"),
        ("PGD epsilon", "0.005, 0.01, 0.02, 0.03, 0.05"),
        ("PGD alpha / steps", "0.005 / 7"),
        ("设备", "NVIDIA GeForce RTX 4070 Ti"),
        ("保存目录", "results/02_attack/fgsm_pgd 与 reports/figures、reports/tables"),
    ]
    add_table(doc, ["配置项", "设置"], rows, [2500, 6860], "表 15  对抗攻击实验配置")

    doc.add_heading("4.4 攻击实验结果", level=2)
    add_table(
        doc,
        ["Attack", "Epsilon", "Clean Acc.", "Adv. Acc.", "Success", "Mean Conf. Drop"],
        attack_rows(attack_df),
        [1100, 1200, 1600, 1600, 1600, 2260],
        "表 16  FGSM 与 PGD 攻击完整结果",
        font_size=8.7,
    )
    add_figure(doc, "fig_11_attack_accuracy_curve.png", "图 14  不同 epsilon 下对抗准确率变化", width=5.95)
    add_figure(doc, "fig_12_attack_success_curve.png", "图 15  不同 epsilon 下攻击成功率变化", width=5.95)
    add_figure(doc, "fig_13_attack_confidence_drop.png", "图 16  不同 epsilon 下平均置信度下降", width=5.95)
    add_para(
        doc,
        "结果显示，epsilon 增大时对抗准确率持续下降，攻击成功率明显上升。在 epsilon=0.03 下，FGSM 将准确率降至 67.10%，PGD 将准确率降至 60.00%。PGD 通过迭代更新和投影搜索，通常比 FGSM 更强；但在 epsilon=0.05 处二者差距缩小，说明大扰动下单步攻击也能产生较强破坏。",
    )

    doc.add_heading("4.5 扰动可感知性：证明这是真正的小扰动攻击", level=2)
    add_para(
        doc,
        "本项目新增扰动可感知性实验，将归一化空间的对抗样本反归一化回像素空间，计算平均 Linf、平均绝对扰动、MSE 和 PSNR。由于 Normalize 使用 std=(0.2672,0.2564,0.2629)，归一化空间 epsilon=0.03 反映到像素空间约为 0.03 * std，即约 0.008，约等于 2/255。人眼看到的 Delta heatmap 是按样本最大扰动重新归一化的热图，目的是展示扰动位置，不代表真实图像变化幅度。",
    )
    add_formula(doc, "x_norm = (x_pixel - mean) / std,    Delta_pixel_c = std_c * Delta_norm_c")
    rows = []
    for _, row in pert_df.sort_values(["attack", "epsilon_normalized_space"]).iterrows():
        rows.append(
            [
                row.attack.upper(),
                num(row.epsilon_normalized_space, 3),
                num(row.mean_linf_pixel_space, 6),
                num(row.mean_abs_delta_pixel_space, 6),
                f"{row.mean_psnr_db:.2f}",
                f"{int(row.successful_changes)}/{int(row.total_images)}",
            ]
        )
    add_table(
        doc,
        ["Attack", "Norm eps", "Mean pixel Linf", "Mean |Delta|", "PSNR(dB)", "Changed"],
        rows,
        [1100, 1200, 1900, 1800, 1500, 1860],
        "表 17  对抗扰动像素空间可感知性指标",
        font_size=8.8,
    )
    add_figure(doc, "fig_20_perturbation_psnr_bar.png", "图 17  不同攻击设置下平均 PSNR，越高表示越不明显", width=5.70)
    add_figure(
        doc,
        "fig_21_fgsm_perceptual_grid.png",
        "图 18  FGSM epsilon=0.03：真实对抗图与放大扰动视图",
        width=6.25,
        note="说明：第二列 Adversarial 是真实对抗样本；Delta x30 与 Delta heat 是人为放大后的分析视图。",
    )
    add_figure(
        doc,
        "fig_22_pgd_perceptual_grid.png",
        "图 19  PGD epsilon=0.03：真实对抗图与放大扰动视图",
        width=6.25,
        note="说明：Delta heat 用于观察扰动集中区域，不是模型输入图像，也不是肉眼真实看到的扰动幅度。",
    )
    add_figure(
        doc,
        "fig_14_fgsm_eps003_triplets.png",
        "图 20  FGSM epsilon=0.03 攻击样例，Noise 列为放大热力图",
        width=6.20,
    )
    add_figure(
        doc,
        "fig_15_pgd_eps003_triplets.png",
        "图 21  PGD epsilon=0.03 攻击样例，Noise 列为放大热力图",
        width=6.20,
    )

    doc.add_heading("4.6 攻击部分对比与消融", level=2)
    fgsm003 = attack_df[(attack_df.attack == "fgsm") & (attack_df.epsilon.round(3) == 0.03)].iloc[0]
    pgd003 = attack_df[(attack_df.attack == "pgd") & (attack_df.epsilon.round(3) == 0.03)].iloc[0]
    rows = [
        ("攻击方法消融", "FGSM vs PGD", f"{pct(fgsm003.adversarial_accuracy)} vs {pct(pgd003.adversarial_accuracy)}", "PGD 更强，说明多步优化更接近局部最坏扰动"),
        ("扰动预算消融", "epsilon 从 0.005 到 0.05", "准确率单调下降趋势明显", "epsilon 控制攻击强度与可感知性"),
        ("可感知性消融", "epsilon=0.01 vs 0.03", "PSNR 约从 52 dB 降至 42-44 dB", "扰动增大但仍属于小幅像素变化"),
        ("展示方式对照", "Adversarial vs Delta heat", "真实图像差异小，热图差异显著", "热图是解释工具，不是攻击本身"),
    ]
    add_table(doc, ["消融类型", "比较项", "现象", "结论"], rows, [1700, 2500, 2500, 2660], "表 18  攻击实验对比与消融结论")


def add_section_4_extended(doc, data):
    doc.add_heading("4.7 补充验证实验：随机噪声对照、PGD 步数消融与置信边界变化", level=2)
    random_df = data["random_noise"]
    step_df = data["pgd_steps"]
    margin_df = data["margin_shift"]
    attack_df = data["attack"]

    add_para(
        doc,
        "为了进一步证明本项目中的攻击是真正意义上的对抗扰动，而不是普通随机噪声造成的图像损坏，本项目增加了随机 Linf 噪声对照实验。对照实验在相同 epsilon 预算内从 [-epsilon, epsilon] 均匀采样随机扰动，并与 FGSM/PGD 的梯度定向扰动比较。若随机噪声几乎不影响模型，而梯度扰动显著降低准确率，则说明攻击利用了模型损失曲面的敏感方向，而不是简单把图片弄脏。",
    )
    rows = []
    fgsm = attack_df[attack_df.attack == "fgsm"].set_index("epsilon")
    pgd = attack_df[attack_df.attack == "pgd"].set_index("epsilon")
    for _, row in random_df.sort_values("epsilon").iterrows():
        eps = float(row.epsilon)
        rows.append(
            [
                num(eps, 3),
                pct(row.perturbed_accuracy),
                pct(fgsm.loc[eps].adversarial_accuracy),
                pct(pgd.loc[eps].adversarial_accuracy),
                pct(row.success_rate_on_clean_correct),
            ]
        )
    add_table(
        doc,
        ["Epsilon", "Random Acc.", "FGSM Acc.", "PGD Acc.", "Random Success"],
        rows,
        [1400, 1900, 1900, 1900, 2260],
        "表 19  随机 Linf 噪声与梯度对抗攻击对照",
        font_size=8.8,
    )
    add_figure(doc, "fig_27_random_vs_adversarial_accuracy.png", "图 22  随机噪声对照与梯度攻击准确率对比", width=6.05)
    add_figure(
        doc,
        "fig_32_random_vs_pgd_visual_grid.png",
        "图 23  随机噪声与 PGD 对抗样本视觉对照",
        width=6.25,
        note="说明：Random 与 PGD 使用相同 epsilon=0.03 扰动预算，但只有 PGD 通过梯度方向显著改变模型预测。",
    )
    rand003 = random_df[np.isclose(random_df.epsilon, 0.03)].iloc[0]
    pgd003 = attack_df[(attack_df.attack == "pgd") & (np.isclose(attack_df.epsilon, 0.03))].iloc[0]
    add_para(
        doc,
        f"结果非常直接：epsilon=0.03 时，随机扰动准确率为 {pct(rand003.perturbed_accuracy)}，几乎不破坏模型；同样预算下 PGD 准确率为 {pct(pgd003.adversarial_accuracy)}。这说明对抗样本不是依靠“明显噪声”取胜，而是沿损失函数梯度方向寻找更靠近错误决策区域的输入。",
    )

    add_para(
        doc,
        "PGD 步数消融用于解释多步攻击为什么比单步攻击强。固定 epsilon=0.03、alpha=0.005，只改变迭代步数。随着步数增加，PGD 在同一扰动球内进行更多次局部搜索，通常能够找到更高损失、更容易误分类的点。",
    )
    rows = []
    for _, row in step_df.sort_values("pgd_steps").iterrows():
        rows.append(
            [
                int(row.pgd_steps),
                pct(row.adversarial_accuracy),
                pct(row.attack_success_rate),
                num(row.mean_confidence_drop, 4),
            ]
        )
    add_table(
        doc,
        ["PGD Steps", "Adv. Acc.", "Attack Success", "Mean Conf. Drop"],
        rows,
        [1800, 2200, 2500, 2860],
        "表 20  PGD 迭代步数消融结果，epsilon=0.03",
        font_size=8.8,
    )
    add_figure(doc, "fig_28_pgd_step_ablation.png", "图 24  PGD 迭代步数对攻击强度的影响", width=5.90)
    last = step_df.sort_values("pgd_steps").iloc[-1]
    add_para(
        doc,
        f"PGD 从 1 步增加到 20 步时，对抗准确率从 {pct(step_df.sort_values('pgd_steps').iloc[0].adversarial_accuracy)} 下降到 {pct(last.adversarial_accuracy)}，攻击成功率升至 {pct(last.attack_success_rate)}。这验证了 PGD 的技术核心并非更大的扰动预算，而是在相同 epsilon 约束内进行更充分的优化。",
    )

    rows = []
    for _, row in margin_df.iterrows():
        rows.append(
            [
                row.condition,
                num(row.mean_top1_top2_margin, 4),
                num(row.median_top1_top2_margin, 4),
                pct(row["low_margin_rate_lt_0.20"]),
            ]
        )
    add_table(
        doc,
        ["Condition", "Mean Margin", "Median Margin", "Low-margin Rate"],
        rows,
        [2600, 2200, 2200, 2360],
        "表 21  Top-1 与 Top-2 概率间隔变化",
        font_size=8.8,
    )
    add_figure(doc, "fig_31_margin_shift_histogram.png", "图 25  Clean、随机扰动和对抗扰动下的概率间隔分布", width=5.95)
    add_para(
        doc,
        "概率间隔用于衡量模型分类决策的稳定性。Clean 与随机扰动的 top-1/top-2 间隔几乎重合，而 FGSM/PGD 会显著压缩间隔，说明对抗扰动不仅改变最终标签，还会把样本推向更不稳定的决策区域。",
    )


def defense_rows(df, eps):
    rows = []
    for attack in ["fgsm", "pgd"]:
        subset = df[(df.attack == attack) & (df.epsilon.round(3) == eps)]
        before = subset.iloc[0].adversarial_accuracy_before_defense
        for _, row in subset.sort_values("defense").iterrows():
            rows.append(
                [
                    attack.upper(),
                    row.defense,
                    pct(before),
                    pct(row.defended_accuracy),
                    pct(row.defended_accuracy - before),
                    pct(row.recovery_rate_on_successful_attacks),
                ]
            )
    return rows


def add_section_5(doc, data):
    doc.add_heading("5. 输入预处理防御原理、实验与消融", level=1)
    defense_df = data["defense"]

    doc.add_heading("5.1 防御思想与数学解释", level=2)
    add_para(
        doc,
        "输入预处理防御不改变模型参数，而是在对抗图像进入模型前做像素空间变换。它的直觉是：许多对抗扰动集中在高频细节或局部像素变化上，适度平滑、滤波或压缩可能削弱扰动方向，使输入重新回到模型较稳定的判别区域。该防御部署成本低，适合作为轻量基线，但不能保证抵御自适应攻击。",
    )
    rows = [
        ("Gaussian Blur", "用高斯核与图像卷积，降低高频变化", "平滑局部噪声，但可能模糊边缘"),
        ("Median Filter", "窗口内取中值，抑制局部异常像素", "对椒盐式异常更有效，对连续扰动有限"),
        ("JPEG Compression", "DCT 变换后量化高频分量，再重建图像", "可丢弃部分高频扰动，实验中效果最好"),
    ]
    add_table(doc, ["防御方法", "原理", "可能影响"], rows, [1900, 3900, 3560], "表 19  输入预处理防御原理")
    add_figure(doc, "fig_26_defense_data_link.png", "图 22  输入预处理防御数据链路", width=6.25)

    doc.add_heading("5.2 防御实验目的、配置与内容", level=2)
    rows = [
        ("实验目的", "评估轻量输入变换能否恢复被 FGSM/PGD 攻击破坏的准确率"),
        ("模型", "同一 ResNet18 best_model.pth，冻结参数"),
        ("测试样本", "使用攻击实验保存的 3000 张 selected_eval_indices"),
        ("攻击设置", "FGSM/PGD，epsilon=0.03 和 0.05；PGD alpha=0.005, steps=7"),
        ("防御设置", "Gaussian kernel=3 sigma=0.6；Median kernel=3；JPEG quality=75"),
        ("评价指标", "defended accuracy、recovery rate on successful attacks、mean defended confidence"),
        ("保存产物", "input_defense_metrics.csv、防御曲线、epsilon=0.03 恢复样例图"),
    ]
    add_table(doc, ["配置项", "设置"], rows, [2500, 6860], "表 20  输入防御实验配置")

    doc.add_heading("5.3 防御结果", level=2)
    add_table(
        doc,
        ["Attack", "Defense", "Before", "Defended", "Acc. Gain", "Recovery"],
        defense_rows(defense_df, 0.03),
        [1050, 2200, 1450, 1600, 1450, 1610],
        "表 21  epsilon=0.03 输入防御结果",
        font_size=8.7,
    )
    add_table(
        doc,
        ["Attack", "Defense", "Before", "Defended", "Acc. Gain", "Recovery"],
        defense_rows(defense_df, 0.05),
        [1050, 2200, 1450, 1600, 1450, 1610],
        "表 22  epsilon=0.05 输入防御结果",
        font_size=8.7,
    )
    add_figure(doc, "fig_16_fgsm_input_defense_curve.png", "图 23  FGSM 攻击下不同输入防御方法的准确率", width=5.95)
    add_figure(doc, "fig_17_pgd_input_defense_curve.png", "图 24  PGD 攻击下不同输入防御方法的准确率", width=5.95)
    add_figure(doc, "fig_18_input_defense_accuracy_bar.png", "图 25  epsilon=0.03 下不同防御方法横向对比", width=6.20)
    add_figure(doc, "fig_19_input_defense_examples.png", "图 26  输入预处理防御恢复样例", width=6.20)
    add_para(
        doc,
        "从结果看，JPEG Compression 在两种攻击和两档 epsilon 下均表现最好。以 PGD epsilon=0.03 为例，攻击后准确率为 60.30%，JPEG 防御后提升至 79.43%，对成功攻击样本的恢复率为 50.84%。Gaussian Blur 与 Median Filter 也有提升，但恢复幅度较小，说明简单平滑并不能充分破坏梯度构造出的有效扰动。",
    )

    doc.add_heading("5.4 防御消融与局限", level=2)
    pgd003 = defense_df[(defense_df.attack == "pgd") & (defense_df.epsilon.round(3) == 0.03)]
    best = pgd003.sort_values("defended_accuracy", ascending=False).iloc[0]
    rows = [
        ("无防御 vs 有防御", "before_defense 与三种变换", f"PGD eps=0.03 从 {pct(pgd003.iloc[0].adversarial_accuracy_before_defense)} 最高恢复到 {pct(best.defended_accuracy)}", "输入变换有效但不能完全恢复 clean accuracy"),
        ("防御方法消融", "Gaussian / Median / JPEG", f"最佳为 {best.defense}", "高频压缩比简单平滑更适合当前扰动"),
        ("扰动强度消融", "epsilon=0.03 vs 0.05", "epsilon=0.05 下恢复后准确率更低", "扰动预算越大，输入防御越难完全恢复"),
        ("部署代价对比", "无需训练 vs 对抗训练", "输入预处理快速可用", "但自适应攻击下可能被绕过"),
    ]
    add_table(doc, ["消融维度", "比较项", "结果", "结论"], rows, [1700, 2500, 2600, 2560], "表 23  输入防御消融结论")
    add_callout(
        doc,
        "防御边界",
        "输入预处理防御是轻量基线，不是完备安全方案。真正强防御应继续验证对抗训练、随机化防御、自适应攻击和更严格的鲁棒评估。本报告只把已完成且有数据支撑的 JPEG/Blur/Median 结果作为正式结论。",
    )


def add_section_5_extended(doc, data):
    doc.add_heading("5.5 防御参数扫描：鲁棒性与干净准确率的权衡", level=2)
    sweep_df = data["defense_sweep"]
    add_para(
        doc,
        "基础防御实验只比较了固定参数的 Gaussian Blur、Median Filter 和 JPEG Compression。为了进一步分析防御强度，本项目增加参数扫描：Gaussian sigma 取 0.3、0.6、1.0；Median kernel 取 3、5；JPEG quality 取 95、75、50。该实验同时评估 clean transformed accuracy 与 PGD epsilon=0.03 defended accuracy，从而观察防御是否以牺牲正常样本识别为代价。",
    )
    rows = []
    selected = sweep_df.sort_values("pgd_defended_accuracy", ascending=False)
    for _, row in selected.iterrows():
        rows.append(
            [
                row.variant,
                pct(row.clean_transformed_accuracy),
                pct(row.pgd_defended_accuracy),
                pct(row.clean_accuracy_drop),
                pct(row.pgd_accuracy_gain),
            ]
        )
    add_table(
        doc,
        ["Variant", "Clean Acc.", "PGD Def. Acc.", "Clean Drop", "PGD Gain"],
        rows,
        [2600, 1700, 1900, 1500, 1660],
        "表 24  输入防御参数扫描结果，按 PGD 防御准确率排序",
        font_size=8.5,
    )
    add_figure(doc, "fig_29_defense_parameter_tradeoff.png", "图 29  输入防御参数扫描：clean accuracy 与 robust accuracy 权衡", width=6.05)
    add_figure(doc, "fig_30_defense_parameter_bar.png", "图 30  PGD epsilon=0.03 下不同防御参数的准确率对比", width=6.05)
    best = selected.iloc[0]
    jpeg75 = sweep_df[sweep_df.variant == "jpeg_q75"].iloc[0]
    add_para(
        doc,
        f"扫描结果显示，JPEG Compression 的防御强度随压缩增强而提高：jpeg_q95、jpeg_q75、jpeg_q50 分别取得不同程度恢复，其中 jpeg_q50 的 PGD 防御准确率最高，为 {pct(best.pgd_defended_accuracy)}，但 clean accuracy 下降到 {pct(best.clean_transformed_accuracy)}。相比之下，jpeg_q75 的鲁棒准确率为 {pct(jpeg75.pgd_defended_accuracy)}，clean accuracy 为 {pct(jpeg75.clean_transformed_accuracy)}，是更均衡的设置。",
    )
    add_para(
        doc,
        "这个结果说明输入防御不是简单地“越强越好”。较强压缩、较大滤波窗口会更有效地削弱对抗扰动，但也会损失交通标志边缘、数字和颜色细节，导致正常样本准确率下降。报告中保留 JPEG quality=75 作为主防御配置，是因为它兼顾 clean 性能和鲁棒恢复；参数扫描则作为更深入的消融证据。",
    )

    doc.add_heading("5.6 PR #6 补充验证：可解释性、JPEG 消融、全测试集与运行效率", level=2)
    add_para(
        doc,
        "在已有输入防御参数扫描基础上，本项目进一步运行 PR #6 中的补充实验。新增实验没有重新训练模型，而是基于同一个 best_model.pth 做推理评估：一是用 Grad-CAM 解释 clean、attack、defense 三阶段模型关注区域；二是单独扫描 JPEG quality=50/75/90；三是在完整测试集上验证关键攻击与防御配置；四是统计推理、攻击生成和防御处理的运行效率。",
    )

    jpeg_df = data["jpeg_quality"]
    eps003 = jpeg_df[(jpeg_df.attack.isin(["fgsm", "pgd"])) & (np.isclose(jpeg_df.epsilon, 0.03))]
    rows = []
    for _, row in eps003.sort_values(["attack", "jpeg_quality"]).iterrows():
        rows.append(
            [
                row.attack.upper(),
                f"Q{int(row.jpeg_quality)}",
                pct(row.accuracy_before_jpeg),
                pct(row.accuracy_after_jpeg),
                pct(row.recovery_rate_on_successful_attacks),
                num(row.mean_psnr_vs_attack, 2),
            ]
        )
    add_table(
        doc,
        ["Attack", "Quality", "Before", "After", "Recovery", "PSNR"],
        rows,
        [1100, 1200, 1500, 1500, 1700, 2360],
        "表 25  JPEG quality 消融结果，epsilon=0.03",
        font_size=8.3,
    )
    add_figure(doc, "fig_36_jpeg_quality_eps003_bar.png", "图 31  epsilon=0.03 下 JPEG quality 防御准确率对比", width=5.95)
    add_figure(doc, "fig_37_jpeg_clean_accuracy_cost.png", "图 32  JPEG quality 对 clean accuracy 的影响", width=5.75)
    add_para(
        doc,
        "JPEG quality 消融表明，Q50 对 FGSM/PGD 的鲁棒恢复最强，但 clean accuracy 下降也更明显；Q90 最接近正常图像，但对抗恢复较弱；Q75 保持了较好的折中，因此主防御实验选择 Q75 具有数据依据。",
    )

    gradcam_df = data["gradcam"]
    rows = []
    for _, row in gradcam_df.iterrows():
        rows.append([int(row.case_id), row.label_name, int(row.clean_pred), int(row.adv_pred), int(row.def_pred), num(row.def_conf, 3)])
    add_table(
        doc,
        ["Case", "True Label", "Clean", "Attack", "Defense", "Def. Conf."],
        rows,
        [800, 2900, 1000, 1000, 1100, 2560],
        "表 26  Grad-CAM 案例预测记录",
        font_size=8.2,
    )
    add_figure(doc, "fig_33_gradcam_clean_attack_defense.png", "图 33  Clean、Attack、Defense 三阶段 Grad-CAM 对比", width=5.95)
    add_para(
        doc,
        "Grad-CAM 结果显示，攻击后热力图常出现偏移或扩散，模型不再稳定关注交通标志主体；JPEG 防御恢复正确预测后，关注区域更接近标志主体。这为“扰动改变模型内部特征利用方式”提供了可解释性证据。",
    )

    full_attack_df = data["full_attack"]
    rows = []
    for _, row in full_attack_df.iterrows():
        rows.append([row.attack.upper(), int(row.total_images), pct(row.clean_accuracy), pct(row.adversarial_accuracy), pct(row.attack_success_rate)])
    add_table(
        doc,
        ["Attack", "Images", "Clean Acc.", "Adv. Acc.", "Success"],
        rows,
        [1200, 1300, 1900, 1900, 3060],
        "表 27  完整测试集攻击验证",
        font_size=8.5,
    )
    full_def_df = data["full_defense"]
    rows = []
    for _, row in full_def_df.iterrows():
        rows.append([row.attack.upper(), int(row.total_images), pct(row.adversarial_accuracy_before_defense), pct(row.defended_accuracy), pct(row.recovery_rate_on_successful_attacks)])
    add_table(
        doc,
        ["Attack", "Images", "Before", "After JPEG Q75", "Recovery"],
        rows,
        [1200, 1300, 1800, 2400, 2660],
        "表 28  完整测试集 JPEG Q75 防御验证",
        font_size=8.5,
    )

    runtime_df = data["runtime"].sort_values("mean_ms_per_image")
    rows = []
    for _, row in runtime_df.iterrows():
        rows.append([row.operation, num(row.mean_ms_per_image, 3), num(row.mean_images_per_second, 1)])
    add_table(doc, ["Operation", "ms/image", "Images/s"], rows, [4300, 2000, 3060], "表 29  运行效率统计", font_size=8.5)
    add_figure(doc, "fig_40_runtime_ms_per_image.png", "图 34  不同操作单图耗时", width=5.95)
    add_figure(doc, "fig_41_runtime_images_per_second.png", "图 35  不同操作吞吐率", width=5.95)
    add_para(
        doc,
        "完整测试集验证说明 3000 张抽样实验趋势不是偶然；运行效率统计则说明 PGD 由于多步反向传播最慢，FGSM 次之，JPEG 防御和普通推理开销较低。该结果可直接支撑答辩 Demo 的流程设计。",
    )

    doc.add_heading("5.7 更严格防御验证：类别鲁棒性、失败案例、自适应攻击与对抗训练", level=2)
    add_para(
        doc,
        "为了避免防御部分只停留在“平均指标提升”，本项目继续补充四类验证：按类别统计 PGD 脆弱性与 JPEG 恢复率，挑出 JPEG 防御失败案例，使用 BPDA 思想构造知道 JPEG 防御存在的自适应攻击，并进行 FGSM 对抗训练微调，比较输入级防御与模型级防御的差异。",
    )

    per_class_df = data["per_class_robustness"].copy()
    pc_valid = per_class_df[per_class_df.clean_correct >= 5]
    vulnerable = pc_valid.sort_values("attack_success_rate_on_clean_correct", ascending=False).head(8)
    rows = []
    for _, row in vulnerable.iterrows():
        rows.append(
            [
                int(row.class_id),
                row.label_name,
                int(row.total_images),
                pct(row.pgd_accuracy),
                pct(row.jpeg_defended_accuracy),
                pct(row.attack_success_rate_on_clean_correct),
                pct(row.recovery_rate_on_successful_attacks),
            ]
        )
    add_table(
        doc,
        ["Class", "Label", "Images", "PGD Acc.", "JPEG Acc.", "Attack Success", "Recovery"],
        rows,
        [750, 2600, 950, 1250, 1250, 1450, 1110],
        "表 30  PGD epsilon=0.03 下最脆弱类别示例",
        font_size=7.7,
    )
    add_figure(doc, "fig_42_per_class_attack_success.png", "图 36  PGD 下攻击成功率最高的类别", width=6.10)
    add_figure(doc, "fig_43_per_class_defense_recovery.png", "图 37  JPEG Q75 恢复率最高的类别", width=6.10)
    add_para(
        doc,
        "类别级结果表明，攻击脆弱性不是均匀分布的。End speed and passing limits、End no passing、End speed limit 80 等类别在 PGD 下更容易被推过决策边界，通常与样本数量少、类别间形状相似或标志主体细节较弱有关；而 Road narrows right、Traffic signals、Stop 等类别在 JPEG 后恢复更明显，说明压缩对不同类别的扰动结构影响不同。",
    )

    failure_df = data["defense_failure_cases"]
    rows = []
    for _, row in failure_df.iterrows():
        rows.append(
            [
                int(row.case_id),
                f"{int(row.label)} {row.label_name}",
                int(row.clean_pred),
                int(row.pgd_pred),
                int(row.jpeg_pred),
                num(row.clean_conf, 3),
                num(row.pgd_conf, 3),
                num(row.jpeg_conf, 3),
            ]
        )
    add_table(
        doc,
        ["Case", "True Label", "Clean", "PGD", "JPEG", "C Conf.", "P Conf.", "J Conf."],
        rows,
        [700, 2500, 800, 800, 800, 1000, 1000, 1760],
        "表 31  JPEG Q75 未恢复的典型失败案例",
        font_size=7.5,
    )
    add_figure(doc, "fig_44_pgd_jpeg_failure_examples.png", "图 38  PGD 攻击后 JPEG Q75 仍失败的样例", width=6.20)
    add_para(
        doc,
        "失败案例用于补足深度反思链路。表中部分样本在 PGD 后已经以很高置信度错分，JPEG 压缩虽然改变了高频细节，但未能把样本拉回正确类别。这说明输入预处理只是在输入空间破坏一部分扰动结构，并没有真正重塑 ResNet18 的决策边界。",
    )

    adaptive = data["adaptive_jpeg"].iloc[0]
    rows = [
        ("Clean", pct(adaptive.clean_accuracy), "正常测试输入"),
        ("Standard PGD before JPEG", pct(adaptive.standard_pgd_accuracy_before_jpeg), "不考虑防御的普通 PGD"),
        ("Standard PGD + JPEG Q75", pct(adaptive.standard_pgd_accuracy_after_jpeg), "先攻击再压缩，模拟非自适应攻击"),
        ("Adaptive BPDA + JPEG Q75", pct(adaptive.adaptive_bpda_accuracy_after_jpeg), "攻击时把 JPEG 近似纳入优化"),
    ]
    add_table(
        doc,
        ["场景", "Accuracy", "含义"],
        rows,
        [3000, 1700, 4660],
        "表 32  JPEG 防御下普通 PGD 与自适应 BPDA 攻击对比",
        font_size=8.2,
    )
    add_figure(doc, "fig_45_adaptive_jpeg_bpda_accuracy.png", "图 39  JPEG 防御在普通 PGD 与自适应攻击下的准确率", width=5.95)
    add_figure(doc, "fig_46_adaptive_jpeg_bpda_examples.png", "图 40  自适应 BPDA 攻击使 JPEG 防御失效的样例", width=6.20)
    add_para(
        doc,
        "JPEG 压缩不可直接求导，因此自适应实验采用 BPDA 思想：前向传播仍执行 JPEG，反向传播用恒等映射近似其梯度。结果显示 BPDA 后 JPEG 防御准确率为 "
        f"{pct(adaptive.adaptive_bpda_accuracy_after_jpeg)}，低于普通 PGD+JPEG 的 {pct(adaptive.standard_pgd_accuracy_after_jpeg)}。下降幅度不大，但它证明防御结论依赖威胁模型：输入预处理能缓解非自适应攻击，却不能被解释为严格安全证明。",
    )

    adv_train_df = data["adv_training_robust"]
    attack_df = data["attack"]
    rows = []
    clean_row = adv_train_df[adv_train_df.attack == "clean"].iloc[0]
    base_clean = attack_df.iloc[0].clean_accuracy
    rows.append(["Clean", "-", pct(base_clean), pct(clean_row.accuracy), pct(clean_row.accuracy - base_clean)])
    for attack_name in ["fgsm", "pgd"]:
        for eps in [0.03, 0.05]:
            base = attack_df[(attack_df.attack == attack_name) & (np.isclose(attack_df.epsilon, eps))].iloc[0]
            adv = adv_train_df[(adv_train_df.attack == attack_name) & (np.isclose(adv_train_df.epsilon, eps))].iloc[0]
            rows.append([attack_name.upper(), num(eps, 2), pct(base.adversarial_accuracy), pct(adv.accuracy), pct(adv.accuracy - base.adversarial_accuracy)])
    add_table(
        doc,
        ["Attack", "Eps", "Baseline Acc.", "Adv Train Acc.", "Gain"],
        rows,
        [1200, 900, 2100, 2300, 2860],
        "表 33  FGSM 对抗训练后的模型级防御效果",
        font_size=8.3,
    )
    add_figure(doc, "fig_47_adv_training_loss_curve.png", "图 41  对抗训练损失曲线", width=5.80)
    add_figure(doc, "fig_48_adv_training_accuracy_curve.png", "图 42  对抗训练准确率曲线", width=5.80)
    add_figure(doc, "fig_49_adv_training_robust_curve.png", "图 43  对抗训练前后鲁棒准确率曲线", width=5.95)
    add_figure(doc, "fig_50_adv_training_eps003_bar.png", "图 44  epsilon=0.03 下 baseline 与对抗训练模型对比", width=5.95)
    add_figure(doc, "fig_51_adv_training_robust_examples.png", "图 45  对抗训练模型抵抗 PGD 的样例", width=6.05)
    add_para(
        doc,
        "对抗训练属于模型级防御：训练时混入 FGSM 对抗样本，使模型在更新参数时直接学习更平滑、更稳定的局部决策边界。本实验从 baseline ResNet18 继续微调 4 个 epoch。结果显示 clean accuracy 从 "
        f"{pct(base_clean)} 到 {pct(clean_row.accuracy)}，没有明显牺牲正常识别；PGD epsilon=0.03 鲁棒准确率从 {pct(attack_df[(attack_df.attack == 'pgd') & (np.isclose(attack_df.epsilon, 0.03))].iloc[0].adversarial_accuracy)} 提升到 {pct(adv_train_df[(adv_train_df.attack == 'pgd') & (np.isclose(adv_train_df.epsilon, 0.03))].iloc[0].accuracy)}。这说明模型级防御比单纯输入预处理更能提升鲁棒性，但训练成本更高，且仍需更强 PGD/TRADES 等训练策略进一步验证。",
    )


def add_section_6(doc):
    doc.add_heading("6. 实现、复现与结果保存", level=1)
    doc.add_heading("6.1 代码结构", level=2)
    rows = [
        ("configs/baseline_resnet18.yaml", "基础识别训练配置"),
        ("configs/attack_fgsm_pgd.yaml", "FGSM/PGD 攻击评估配置"),
        ("configs/defense_input_preprocessing.yaml", "输入预处理防御配置"),
        ("src/models/classifiers.py", "SimpleCNN 与 ResNet18 构建函数，小图像适配逻辑"),
        ("src/attacks/methods.py", "FGSM 与 PGD 攻击函数"),
        ("src/evaluate_attacks.py", "攻击评估、攻击曲线和对抗样例图生成"),
        ("src/evaluate_input_defense.py", "输入预处理防御评估与恢复样例图生成"),
        ("src/analyze_perturbations.py", "像素空间 Linf、MSE、PSNR 与 Delta x30 图生成"),
        ("src/extended_experiments.py", "随机噪声对照、PGD 步数消融、防御参数扫描与概率间隔分析"),
        ("src/visualization/gradcam_analysis.py", "Grad-CAM clean/attack/defense 可解释性分析"),
        ("src/evaluate_jpeg_ablation.py", "JPEG quality 参数消融实验"),
        ("src/benchmark_runtime.py", "推理、攻击生成和输入防御运行效率统计"),
        ("src/evaluate_per_class_failure.py", "按类别鲁棒性统计与防御失败案例分析"),
        ("src/evaluate_adaptive_jpeg_attack.py", "JPEG 防御下 BPDA 自适应攻击验证"),
        ("src/train_adversarial.py", "FGSM 对抗训练模型级防御实验"),
        ("scripts/build_deep_word_report.py", "本深度实验报告生成脚本"),
    ]
    add_table(doc, ["文件", "功能"], rows, [3400, 5960], "表 24  核心实现文件")

    doc.add_heading("6.2 复现实验命令", level=2)
    commands = [
        "python -m src.data.check_gtsrb --config configs/dataset_check.yaml",
        "python -m src.train_classifier --config configs/baseline_resnet18.yaml",
        "python -m src.evaluate_attacks --config configs/attack_fgsm_pgd.yaml",
        "python -m src.analyze_perturbations",
        "python -m src.evaluate_input_defense --config configs/defense_input_preprocessing.yaml",
        "python -m src.extended_experiments",
        "python -m src.visualization.gradcam_analysis --config configs/explainability_gradcam.yaml",
        "python -m src.evaluate_jpeg_ablation --config configs/defense_jpeg_ablation.yaml",
        "python -m src.benchmark_runtime --config configs/runtime_benchmark.yaml",
        "python -m src.evaluate_attacks --config configs/attack_fgsm_pgd_full_test.yaml",
        "python -m src.evaluate_input_defense --config configs/defense_input_preprocessing_full_test.yaml",
        "python -m src.evaluate_per_class_failure --config configs/per_class_failure_analysis.yaml",
        "python -m src.evaluate_adaptive_jpeg_attack --config configs/adaptive_jpeg_attack.yaml",
        "python -m src.train_adversarial --config configs/defense_adversarial_training.yaml",
        "python scripts/build_deep_word_report.py",
    ]
    for cmd in commands:
        add_code_line(doc, cmd)

    doc.add_heading("6.3 报告使用的主要结果产物", level=2)
    rows = [
        ("数据检查", "reports/tables/table_01_dataset_summary.csv", "reports/figures/fig_01-04"),
        ("基础识别", "table_03_baseline_test_metrics.json、table_04_per_class_metrics.csv", "fig_05-10"),
        ("攻击实验", "table_06_attack_metrics.csv", "fig_11-15"),
        ("扰动分析", "table_08_perturbation_perceptibility_metrics.csv", "fig_20-22"),
        ("防御实验", "table_07_input_defense_metrics.csv", "fig_16-19"),
        ("补充验证", "table_09-12 随机噪声、PGD 步数、防御扫描、margin 指标", "fig_27-32"),
        ("PR #6 补充实验", "table_13-15 Grad-CAM、JPEG quality、runtime", "fig_33-41"),
        ("严格防御验证", "table_16-20 类别鲁棒性、失败案例、自适应攻击、对抗训练", "fig_42-51"),
        ("全测试集验证", "results/02_attack/fgsm_pgd_full_test 与 results/03_defense/input_preprocessing_full_test", "完整 12630 张测试图关键配置验证"),
        ("模型权重", "results/01_baseline/resnet18/checkpoints/best_model.pth", "用于攻击与防御评估"),
        ("对抗训练权重", "results/03_defense/adversarial_training/checkpoints/best_model.pth", "模型级防御评估"),
    ]
    add_table(doc, ["实验模块", "表格/指标产物", "图片/模型产物"], rows, [1800, 4200, 3360], "表 25  结果保存与报告素材索引")


def add_section_7(doc):
    doc.add_heading("7. 已补强后的不足与后续计划", level=1)
    add_para(
        doc,
        "在新增 per-class 鲁棒性、失败案例、自适应 BPDA 攻击和 FGSM 对抗训练后，项目已经从“完整课程实验”进一步接近鲁棒性研究。但仍有一些更高阶内容没有展开，报告中不虚构未完成指标，而把它们作为后续计划。",
    )
    rows = [
        ("模型结构消融", "SimpleCNN、原始 ResNet18 stem、当前适配 ResNet18", "未完成训练", "证明结构选择不是拍脑袋"),
        ("数据增强消融", "无增强 vs 当前增强", "未完成训练", "分析 clean 泛化与鲁棒性的关系"),
        ("更强鲁棒训练", "PGD adversarial training、TRADES、MART", "未完成", "检验 FGSM 训练是否存在梯度遮蔽或过拟合"),
        ("更强自适应攻击", "EOT、多随机重启、更多 BPDA 近似", "已做 BPDA 初版", "更严格评估输入预处理防御"),
        ("可解释性扩展", "Grad-CAM 已完成，可继续加入失败样本 CAM 或多层 CAM", "部分完成", "增强失败原因分析"),
        ("交互 Demo", "上传图片 -> 识别 -> 攻击 -> 防御恢复", "脚本入口未正式录屏", "提高课程展示效果"),
        ("更大规模防御验证", "已完成关键配置全测试集验证，可继续跨模型重复", "部分完成", "验证结论是否跨模型稳定"),
    ]
    add_table(doc, ["待补实验", "目的", "当前状态", "补强价值"], rows, [1800, 3100, 1700, 2760], "表 26  当前仍缺实验与验证计划")


def add_conclusion_roles_refs(doc):
    doc.add_heading("8. 创新点、结论与分工", level=1)
    doc.add_heading("8.1 项目创新点", level=2)
    for item in [
        "从普通交通标志分类扩展到安全性评估，形成识别、攻击、防御闭环。",
        "将 epsilon 从归一化空间解释到像素空间，并用 PSNR、Linf 证明扰动确实较小。",
        "同时比较 FGSM 与 PGD，进行攻击方法与扰动预算双重消融。",
        "加入随机 Linf 噪声对照，证明对抗攻击的有效性来自梯度定向而不是普通噪声。",
        "加入 PGD 步数消融和概率间隔分析，解释迭代优化如何压缩模型决策边界。",
        "进行输入预处理防御消融，比较无防御、Gaussian Blur、Median Filter、JPEG Compression。",
        "进一步扫描防御参数，分析 robust accuracy 与 clean accuracy 的实际权衡。",
        "补充 Grad-CAM 可解释性、JPEG quality 消融、完整测试集关键配置验证和运行效率统计。",
        "新增按类别鲁棒性和防御失败案例，说明平均指标之外的类别差异和失败边界。",
        "新增 BPDA 自适应攻击，检验 JPEG 输入防御在更强威胁模型下的有效性。",
        "新增 FGSM 对抗训练，形成输入级防御和模型级防御的对比。",
        "保存训练曲线、混淆矩阵、对抗样本、放大扰动图、防御恢复案例等中间产物，便于报告和答辩展示。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("8.2 总结", level=2)
    add_para(
        doc,
        "本项目首先建立了 ResNet18 交通标志识别基线，clean 测试准确率达到 97.81%，说明基础识别能力成立。随后在同一模型上进行白盒无目标 FGSM/PGD 攻击，结果显示即使像素空间平均 Linf 约为 0.008 的小扰动，也能显著降低模型准确率。PGD epsilon=0.03 将准确率降至 60.00%，证明模型在决策边界附近存在明显脆弱性。",
    )
    add_para(
        doc,
        "防御实验表明，输入预处理方法能够部分恢复模型性能，其中 JPEG Compression 在当前设置下最有效。进一步参数扫描和 JPEG quality 消融显示，更强压缩通常带来更高鲁棒恢复，但也会损失 clean accuracy，因此防御配置需要在正常识别与鲁棒性之间折中。新增类别级分析和失败案例说明，防御收益并不均匀；BPDA 自适应攻击进一步提醒输入预处理不能被当成严格安全证明。FGSM 对抗训练则提供了模型级防御证据，使 PGD epsilon=0.03 鲁棒准确率提升到 88.63%。整体来看，本项目已经形成较完整的深度学习识别到对抗鲁棒性的实验链路，技术性和展示性均强于单纯复刻分类模型。",
    )

    doc.add_heading("8.3 小组分工模板", level=2)
    rows = [
        ("成员 A", "数据集处理、数据可视化、类别统计", "20%"),
        ("成员 B", "ResNet18 模型结构、训练与 clean 评估", "20%"),
        ("成员 C", "FGSM/PGD 攻击实现、攻击曲线与样例保存", "20%"),
        ("成员 D", "输入防御实现、防御消融与恢复样例", "20%"),
        ("成员 E", "报告、PPT、实验材料整理与答辩展示", "20%"),
    ]
    add_table(doc, ["成员", "工作内容", "贡献占比"], rows, [1600, 6000, 1760], "表 27  小组分工模板，提交前按实际成员修改")

    doc.add_heading("参考文献", level=1)
    refs = [
        "Stallkamp J., Schlipsing M., Salmen J., Igel C. The German Traffic Sign Recognition Benchmark: A multi-class classification competition. IJCNN, 2011.",
        "He K., Zhang X., Ren S., Sun J. Deep Residual Learning for Image Recognition. CVPR, 2016.",
        "Szegedy C. et al. Intriguing properties of neural networks. ICLR, 2014.",
        "Goodfellow I., Shlens J., Szegedy C. Explaining and Harnessing Adversarial Examples. ICLR, 2015.",
        "Madry A., Makelov A., Schmidt L., Tsipras D., Vladu A. Towards Deep Learning Models Resistant to Adversarial Attacks. ICLR, 2018.",
    ]
    for ref in refs:
        add_numbered(doc, ref)


def load_data():
    return {
        "dataset": load_json(TABLE_DIR / "table_02_dataset_overview.json"),
        "baseline": load_json(TABLE_DIR / "table_03_baseline_test_metrics.json"),
        "per_class": pd.read_csv(TABLE_DIR / "table_04_baseline_per_class_metrics.csv"),
        "train_log": pd.read_csv(TABLE_DIR / "table_05_baseline_train_log.csv"),
        "attack": pd.read_csv(TABLE_DIR / "table_06_attack_metrics.csv"),
        "defense": pd.read_csv(TABLE_DIR / "table_07_input_defense_metrics.csv"),
        "perturbation": pd.read_csv(TABLE_DIR / "table_08_perturbation_perceptibility_metrics.csv"),
        "random_noise": pd.read_csv(TABLE_DIR / "table_09_random_noise_control.csv"),
        "pgd_steps": pd.read_csv(TABLE_DIR / "table_10_pgd_step_ablation.csv"),
        "defense_sweep": pd.read_csv(TABLE_DIR / "table_11_defense_parameter_sweep.csv"),
        "margin_shift": pd.read_csv(TABLE_DIR / "table_12_margin_shift_metrics.csv"),
        "gradcam": pd.read_csv(TABLE_DIR / "table_13_gradcam_cases.csv"),
        "jpeg_quality": pd.read_csv(TABLE_DIR / "table_14_jpeg_quality_metrics.csv"),
        "runtime": pd.read_csv(TABLE_DIR / "table_15_runtime_summary.csv"),
        "per_class_robustness": pd.read_csv(TABLE_DIR / "table_16_per_class_robustness.csv"),
        "defense_failure_cases": pd.read_csv(TABLE_DIR / "table_17_defense_failure_cases.csv"),
        "adaptive_jpeg": pd.read_csv(TABLE_DIR / "table_18_adaptive_jpeg_bpda_metrics.csv"),
        "adv_training_log": pd.read_csv(TABLE_DIR / "table_19_adv_training_train_log.csv"),
        "adv_training_robust": pd.read_csv(TABLE_DIR / "table_20_adv_training_robust_metrics.csv"),
        "full_attack": pd.read_csv(ROOT / "results" / "02_attack" / "fgsm_pgd_full_test" / "metrics" / "attack_metrics.csv"),
        "full_defense": pd.read_csv(ROOT / "results" / "03_defense" / "input_preprocessing_full_test" / "metrics" / "input_defense_metrics.csv"),
    }


def build_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generate_diagrams()
    data = load_data()
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_abstract_and_summary(doc, data)
    add_section_1(doc)
    add_section_2(doc, data)
    add_section_3(doc, data)
    add_section_4(doc, data)
    add_section_4_extended(doc, data)
    add_section_5(doc, data)
    add_section_5_extended(doc, data)
    add_section_6(doc)
    add_section_7(doc)
    add_conclusion_roles_refs(doc)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_docx())
