from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from build_deep_word_report import (
    FIG_DIR,
    TABLE_DIR,
    add_callout,
    add_code_line,
    add_figure,
    add_formula,
    add_para,
    add_table,
    generate_diagrams,
    load_json,
    num,
    pct,
    set_run_font,
    setup_document,
)


ROOT_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT_DIR / "reports" / "word"
OUT_PATH = OUT_DIR / "教师要求版-基于ResNet交通标志识别的对抗扰动攻击与输入防御实验报告.docx"

TABLE_NO = 1
FIG_NO = 1


def table(doc, title, headers, rows, widths, explanation, font_size=8.8):
    global TABLE_NO
    add_table(doc, headers, rows, widths, f"表 {TABLE_NO}  {title}", font_size=font_size)
    explain(doc, f"表 {TABLE_NO} 解释", explanation)
    TABLE_NO += 1


def figure(doc, filename, title, explanation, width=6.05, note=None):
    global FIG_NO
    add_figure(doc, filename, f"图 {FIG_NO}  {title}", width=width, note=note)
    explain(doc, f"图 {FIG_NO} 解释", explanation)
    FIG_NO += 1


def explain(doc, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(f"{title}：")
    set_run_font(r, size=10.2, bold=True, color=RGBColor(31, 77, 120))
    r2 = p.add_run(body)
    set_run_font(r2, size=10.2, color=RGBColor(30, 30, 30))


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def load_data():
    return {
        "dataset": load_json(TABLE_DIR / "table_02_dataset_overview.json"),
        "baseline": load_json(TABLE_DIR / "table_03_baseline_test_metrics.json"),
        "per_class": pd.read_csv(TABLE_DIR / "table_04_baseline_per_class_metrics.csv"),
        "train_log": pd.read_csv(TABLE_DIR / "table_05_baseline_train_log.csv"),
        "attack": pd.read_csv(TABLE_DIR / "table_06_attack_metrics.csv"),
        "defense": pd.read_csv(TABLE_DIR / "table_07_input_defense_metrics.csv"),
        "perturbation": pd.read_csv(TABLE_DIR / "table_08_perturbation_perceptibility_metrics.csv"),
        "random_noise": pd.read_csv(TABLE_DIR / "table_09_random_noise_control.csv"),
        "pgd_steps": pd.read_csv(TABLE_DIR / "table_10_pgd_step_ablation.csv"),
        "defense_sweep": pd.read_csv(TABLE_DIR / "table_11_defense_parameter_sweep.csv"),
        "margin_shift": pd.read_csv(TABLE_DIR / "table_12_margin_shift_metrics.csv"),
        "gradcam": pd.read_csv(TABLE_DIR / "table_13_gradcam_cases.csv"),
        "jpeg_quality": pd.read_csv(TABLE_DIR / "table_14_jpeg_quality_metrics.csv"),
        "runtime": pd.read_csv(TABLE_DIR / "table_15_runtime_summary.csv"),
        "full_attack": pd.read_csv(ROOT_DIR / "results" / "02_attack" / "fgsm_pgd_full_test" / "metrics" / "attack_metrics.csv"),
        "full_defense": pd.read_csv(ROOT_DIR / "results" / "03_defense" / "input_preprocessing_full_test" / "metrics" / "input_defense_metrics.csv"),
    }


def add_cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("汇报题目：基于 ResNet 交通标志识别的对抗扰动攻击与输入防御系统设计与实现")
    set_run_font(r, size=19, bold=True, color=RGBColor(46, 116, 181))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("从深度学习识别原理到 FGSM/PGD 攻击、扰动可感知性验证与防御消融")
    set_run_font(r, size=12, color=RGBColor(90, 90, 90))

    rows = [
        ("课程报告类型", "期末大作业汇报报告，按老师六部分模板组织"),
        ("项目对象", "GTSRB 交通标志 43 类图像分类"),
        ("核心算法", "ResNet18 小尺寸图像适配版 + FGSM/PGD 对抗攻击 + 输入预处理防御"),
        ("汇报重点", "不仅说明做了什么，更解释为什么这样做、算法原理是什么、图表结果说明什么"),
        ("提交日期", "2026 年 6 月"),
    ]
    table(
        doc,
        "报告基本信息",
        ["项目", "内容"],
        rows,
        [2200, 7160],
        "本表对应老师模板中的“汇报题目”和项目定位。YOLO 示例是目标检测方向，本项目实际方向是交通标志识别与对抗鲁棒性，因此报告保留老师要求的六部分结构，但算法内容替换为 ResNet 分类、FGSM/PGD 攻击和输入防御。",
    )
    doc.add_page_break()


def part1(doc, data):
    h1(doc, "第一部分：项目背景与任务定义")
    h2(doc, "1.1 应用场景需求")
    add_para(
        doc,
        "交通标志识别是自动驾驶、辅助驾驶和智能交通系统中的关键视觉感知任务。车辆需要识别限速、禁止通行、让行、施工、转向等道路标志，并把识别结果交给后续决策模块。如果交通标志识别出错，车辆可能出现错误限速、错误路径选择或风险提示失败。因此，本项目选择交通标志识别作为基础任务，并进一步研究模型在对抗扰动下的安全性。",
    )
    add_para(
        doc,
        "与普通分类任务不同，本项目关注的是“模型在正常样本上准确”是否足够。深度学习模型可能在 clean test 上表现很好，但面对人眼难以察觉的小扰动时预测错误。对于智能交通场景，这类鲁棒性问题比单纯准确率更值得分析。",
    )

    h2(doc, "1.2 技术挑战：准确性、鲁棒性与部署代价的权衡")
    rows = [
        ("准确性 Accuracy", "模型在正常测试集上的分类正确率", "基础识别必须先达到较高水平，否则攻击实验没有意义"),
        ("鲁棒性 Robustness", "模型面对扰动、压缩、光照变化或恶意攻击时保持正确的能力", "本项目核心研究点，不只看 clean accuracy"),
        ("扰动不可感知性", "扰动在像素空间很小，人眼通常难以直接观察", "需要用 Linf、MSE、PSNR 等指标证明"),
        ("防御代价", "输入防御可能提高对抗样本准确率，但也可能损失正常样本细节", "需要做 clean-robust trade-off 分析"),
    ]
    table(
        doc,
        "项目技术挑战与为什么需要补充实验",
        ["挑战", "含义", "为什么重要"],
        rows,
        [2100, 3600, 3660],
        "本表解释老师要求中的“技术挑战”。对于 YOLO 项目，挑战通常是实时性与准确率；对于本项目，主要挑战是 clean accuracy 与 adversarial robustness 的权衡，以及防御强度与正常识别损失之间的 trade-off。",
    )

    h2(doc, "1.3 任务定义与总流程")
    rows = [
        ("输入", "64 x 64 RGB 交通标志图像", "原始图片统一缩放并归一化"),
        ("输出", "43 类交通标志类别编号和概率", "ResNet18 输出 43 维 logits，再经 softmax 得到概率"),
        ("基础模型", "ResNet18", "作为识别主干网络，先建立 clean baseline"),
        ("攻击实验", "FGSM、PGD", "在 Linf 约束内构造对抗样本"),
        ("防御实验", "Gaussian Blur、Median Filter、JPEG Compression", "不重新训练模型，评估输入预处理能否恢复准确率"),
        ("补充验证", "随机噪声对照、PGD 步数消融、防御参数扫描", "证明攻击不是普通噪声，并分析防御权衡"),
    ]
    table(
        doc,
        "任务定义",
        ["模块", "定义", "报告中作用"],
        rows,
        [1700, 3600, 4060],
        "本表给出项目的输入、输出和实验模块。后续所有实验都围绕同一 ResNet18 模型展开，保证 clean、attack、defense 三类结果可以直接比较。",
    )
    figure(
        doc,
        "fig_23_experiment_pipeline.png",
        "项目总流程：识别、攻击、扰动分析与防御",
        "该图展示完整数据链路：GTSRB 原始图像先经过 resize、augmentation 和 normalization，再进入 ResNet18；模型输出 clean metrics 后，同一模型继续用于 FGSM/PGD 攻击、扰动可感知性分析和输入防御。箭头表示数据流方向，说明本项目不是孤立实验，而是同一模型上的连续鲁棒性评估。",
        width=6.25,
    )

    h2(doc, "1.4 小组分工")
    rows = [
        ("成员 A", "数据集整理、类别统计、样例图和数据分析", "20%"),
        ("成员 B", "ResNet18 模型结构、训练、clean 指标分析", "20%"),
        ("成员 C", "FGSM/PGD 攻击实现、扰动可感知性分析", "20%"),
        ("成员 D", "输入防御、参数扫描、消融实验", "20%"),
        ("成员 E", "报告、PPT、演示材料和答辩组织", "20%"),
    ]
    table(
        doc,
        "小组分工",
        ["成员", "主要职责", "贡献百分比"],
        rows,
        [1500, 6100, 1760],
        "本表对应老师模板中的“小组分工”。提交前可将成员 A-E 替换为真实姓名；贡献百分比总和为 100%。",
    )


def part2(doc, data):
    h1(doc, "第二部分：算法原理深度解析")
    h2(doc, "2.1 深度学习分类识别原理")
    add_para(
        doc,
        "交通标志识别属于监督分类任务。模型输入图像张量 x，输出 43 个类别的 logits。logits 是 softmax 之前的原始分数，不是概率；softmax 会把 logits 转换为各类别概率；交叉熵损失 Cross Entropy 会惩罚真实类别概率过低的情况。",
    )
    add_formula(doc, "z = f_theta(x),    p_i = exp(z_i) / sum_j exp(z_j),    L(x, y; theta) = -log p_y")
    rows = [
        ("logits", "模型最后一层输出的原始类别分数", "越大表示模型越倾向该类别，但还不是概率"),
        ("softmax", "将 logits 转成概率分布的函数", "用于得到 top-1 类别和置信度 confidence"),
        ("Cross Entropy", "分类任务常用损失函数", "训练时最小化，攻击时反过来最大化"),
        ("Accuracy", "正确样本数 / 总样本数", "衡量总体分类正确率"),
        ("Precision/Recall/F1", "精确率、召回率和综合指标", "macro 表示先按类别计算再平均，避免只看大类"),
    ]
    table(
        doc,
        "分类模型核心术语解释",
        ["术语", "中文解释", "本项目中的作用"],
        rows,
        [2200, 3300, 3860],
        "本表解释后续报告中反复出现的英文术语。老师要求参数和单词要解释，因此这里先给出基础分类词汇，避免只堆英文指标。",
    )

    h2(doc, "2.2 ResNet18 模型结构拆解")
    add_para(
        doc,
        "ResNet 的核心思想是残差学习。普通深层网络随着层数增加可能更难优化，ResNet 通过 shortcut connection 让网络学习残差 F(x)，最终输出 y=F(x)+x。这样梯度可以沿恒等路径更稳定地传播，缓解深层网络训练退化问题。",
    )
    add_formula(doc, "Residual Block:    y = F(x, W) + x")
    rows = [
        ("Backbone 骨干网络", "ResNet18", "逐层提取边缘、颜色、形状和语义特征"),
        ("Conv 卷积层", "局部窗口共享权重", "识别交通标志边缘、数字、圆形/三角形边框"),
        ("Receptive Field 感受野", "一个神经元能看到的输入区域", "层数越深，感受野越大，可理解整个标志形状"),
        ("Downsampling 下采样", "降低特征图空间分辨率", "扩大感受野，但过早下采样会丢失小图细节"),
        ("Residual Shortcut", "输入 x 直接加到残差输出上", "让深层网络更容易训练"),
        ("FC 分类层", "Linear(512, 43)", "将高层特征映射到 43 类交通标志"),
    ]
    table(
        doc,
        "ResNet18 结构术语解释",
        ["结构/参数", "中文解释", "为什么这样做"],
        rows,
        [2200, 3300, 3860],
        "本表按照老师模板中 Backbone、特征提取、感受野等要求解释 ResNet。虽然本项目不是 YOLO 检测任务，没有 Neck/FPN/PAN，但 ResNet18 本身就是分类任务的特征提取 Backbone。",
    )
    figure(
        doc,
        "fig_24_resnet_recognition_flow.png",
        "ResNet18 识别流程与残差结构",
        "该图从输入张量开始，依次展示 3x3 卷积、残差阶段、全局平均池化、全连接层和 softmax。图中 y=F(x,W)+x 是残差块核心公式。由于 GTSRB 图像较小，本项目把原始 ResNet18 的 7x7 大卷积和 maxpool 改成 3x3 卷积与 Identity，以避免过早丢失交通标志边缘和数字细节。",
        width=6.25,
    )

    h2(doc, "2.3 对抗扰动攻击原理")
    add_para(
        doc,
        "对抗样本是对原始图像加入小扰动后得到的样本。攻击目标不是让图像在人眼看来完全改变，而是在扰动幅度受限的情况下最大化模型损失，使模型输出错误类别。本项目采用白盒无目标攻击：白盒表示知道模型结构和权重，无目标表示只要求预测错，不指定错成哪一类。",
    )
    add_formula(doc, "x_adv = arg max_{x' in B_inf(x, epsilon)} L(f_theta(x'), y)")
    rows = [
        ("epsilon", "扰动预算，限制最大变化幅度", "epsilon 越大，攻击越强，但可见性也可能增加"),
        ("Linf", "最大范数约束", "限制每个像素通道的最大变化，而不是平均变化"),
        ("FGSM", "Fast Gradient Sign Method，单步梯度符号攻击", "速度快，用 sign(grad_x L) 一步生成扰动"),
        ("PGD", "Projected Gradient Descent，多步投影梯度攻击", "多次更新并投影回 epsilon 球内，通常更强"),
        ("alpha", "PGD 每一步的步长", "控制每次沿梯度方向走多远"),
        ("steps", "PGD 迭代次数", "步数越多，搜索越充分，计算越慢"),
        ("PSNR", "Peak Signal-to-Noise Ratio，峰值信噪比", "越高表示扰动越不明显"),
    ]
    table(
        doc,
        "攻击参数与英文术语解释",
        ["参数/术语", "中文解释", "本项目含义"],
        rows,
        [1900, 3600, 3860],
        "本表解释攻击实验中的核心参数。特别注意：epsilon 是在归一化 tensor 空间设置的，反归一化到像素空间后 epsilon=0.03 约为 2/255 级别，因此实际扰动很小。",
    )
    h2(doc, "2.3.1 两种攻击方式的机制对比：FGSM 与 PGD")
    add_para(
        doc,
        "本项目不是只做一种攻击，而是选择 FGSM 和 PGD 两种代表性白盒攻击。选择 FGSM 是因为它是最经典、最容易解释的单步梯度攻击，可以作为快速攻击基线；选择 PGD 是因为它通过多步迭代近似求解 epsilon 约束下的最坏扰动，通常被视为更强的鲁棒性评估基线。两者的区别不在于是否“加噪声”，而在于如何利用模型梯度搜索让损失增大的扰动方向。",
    )
    add_formula(doc, "FGSM: x_adv = Clip(x + epsilon * sign(grad_x L(f_theta(x), y)))")
    add_formula(doc, "PGD: x_{t+1} = Proj_{B_inf(x,epsilon)}(x_t + alpha * sign(grad_x L(f_theta(x_t), y)))")
    rows = [
        ("FGSM", "Fast Gradient Sign Method，快速梯度符号法", "单步", "只在原图 x 上计算一次输入梯度", "速度快、实现简单，适合做基础攻击基线", "搜索不充分，容易低估模型脆弱性"),
        ("PGD", "Projected Gradient Descent，投影梯度下降/上升攻击", "多步", "每一步更新后投影回 Linf epsilon 球", "攻击更强，更接近局部最坏情况", "计算更慢，需要设置 alpha 和 steps"),
    ]
    table(
        doc,
        "FGSM 与 PGD 攻击机制对比",
        ["攻击", "中文含义", "更新方式", "核心逻辑", "优点", "局限"],
        rows,
        [900, 2100, 1000, 2100, 1600, 1660],
        "本表补充两种攻击方式的原理解读。FGSM 依赖一次线性近似，PGD 则在同一扰动预算内反复寻找更强方向。因此当 epsilon 相同时，PGD 通常比 FGSM 更能暴露模型鲁棒性问题。",
        font_size=7.8,
    )
    rows = [
        ("epsilon", "FGSM 的总扰动幅度，也是 PGD 的扰动球半径", "控制攻击预算；本项目主要解释 eps=0.03 和 0.05"),
        ("alpha", "PGD 每一步的更新步长", "如果 alpha 太小，搜索慢；太大，可能在边界附近震荡"),
        ("steps", "PGD 迭代次数", "步数越多，搜索越充分，但计算成本越高"),
        ("projection", "投影操作", "保证 PGD 每次更新后仍满足 Linf 约束，不让攻击变成无限制改图"),
        ("Clip", "裁剪到合法像素范围", "保证对抗样本仍对应合法图像输入"),
    ]
    table(
        doc,
        "攻击公式中关键参数解释",
        ["参数/操作", "解释", "为什么重要"],
        rows,
        [1900, 3400, 4060],
        "本表专门解释公式中的参数。这样报告不仅展示攻击结果，也说明为什么 PGD 需要 alpha、steps 和 projection，而 FGSM 只需要 epsilon。",
    )
    figure(
        doc,
        "fig_25_attack_threat_model.png",
        "白盒无目标对抗攻击威胁模型",
        "该图解释攻击如何产生：原图 x 前向传播得到损失 L，再对输入求梯度 grad_x L，沿梯度符号方向生成扰动，并投影回 Linf epsilon ball。图中的 Delta heatmap 是放大后的分析视图，不是模型实际输入。",
        width=6.25,
    )

    h2(doc, "2.4 输入预处理防御原理")
    rows = [
        ("Gaussian Blur", "高斯模糊", "用平滑核削弱局部高频扰动，但可能模糊边缘"),
        ("Median Filter", "中值滤波", "用邻域中位数替换像素，抑制局部异常点"),
        ("JPEG Compression", "JPEG 有损压缩", "通过 DCT 量化丢弃部分高频信息，可能削弱对抗扰动"),
        ("Clean-Robust Trade-off", "干净准确率与鲁棒准确率权衡", "防御越强，可能越损失正常图像细节"),
    ]
    table(
        doc,
        "防御方法术语解释",
        ["方法/术语", "中文解释", "为什么可能有效"],
        rows,
        [2300, 3000, 4060],
        "本表说明防御不是重新训练模型，而是在输入进入 ResNet18 前做像素空间变换。它的优点是部署简单，缺点是面对自适应攻击时可能被绕过。",
    )
    h2(doc, "2.4.1 三种输入防御的原理细化")
    add_para(
        doc,
        "输入预处理防御的共同思想是：对抗扰动虽然幅度小，但往往包含模型敏感的高频或局部方向变化。若在送入模型前对图像做空间平滑、局部统计替换或有损压缩，可能破坏这部分扰动结构，使样本重新落回模型较稳定的分类区域。不过，这类方法也可能同时损伤交通标志的边缘、数字和纹理，因此必须分析防御收益和 clean accuracy 损失。",
    )
    rows = [
        ("Gaussian Blur", "空间域线性平滑", "使用高斯核对邻域像素加权平均", "削弱连续高频噪声和细小纹理变化", "sigma 越大越平滑，边缘越容易变糊"),
        ("Median Filter", "空间域非线性滤波", "用窗口内中位数替换中心像素", "抑制局部突变像素，保留部分边缘", "kernel 越大，去噪更强，但小数字和细线可能被抹掉"),
        ("JPEG Compression", "频域有损压缩", "DCT 变换后量化高频系数，再重建图像", "丢弃部分高频扰动，对当前攻击恢复明显", "quality 越低压缩越强，正常图像细节损失越大"),
    ]
    table(
        doc,
        "三种输入防御的机制、参数和局限",
        ["防御", "处理域", "工作机制", "为什么能防御", "参数影响/局限"],
        rows,
        [1350, 1500, 2600, 2300, 1610],
        "本表补充防御原理解读。Gaussian 和 Median 属于空间域处理，直接改变像素邻域；JPEG 属于频域压缩，主要影响高频信息。它们都不是“修复模型”，而是改变输入分布。",
        font_size=7.8,
    )
    add_formula(doc, "Gaussian Blur: I'(u,v) = sum_{i,j} G_sigma(i,j) * I(u-i, v-j)")
    add_formula(doc, "Median Filter: I'(u,v) = median{ I(u+i, v+j) | (i,j) in local window }")
    add_formula(doc, "JPEG: image -> block DCT -> quantization -> inverse DCT -> compressed image")
    add_para(
        doc,
        "从原理上看，Gaussian Blur 更适合处理连续平滑噪声，Median Filter 更适合处理局部异常点，JPEG Compression 更适合去除高频细节。对抗扰动不一定完全等同于自然噪声，因此三种方法的效果必须通过实验比较，不能只凭直觉判断。",
    )
    figure(
        doc,
        "fig_26_defense_data_link.png",
        "输入预处理防御数据链路",
        "该图展示防御链路：对抗图像先反归一化回像素空间，再执行 Blur/Median/JPEG 变换，之后重新归一化并输入同一个冻结 ResNet18。这样可以保证防御评估只改变输入，不改变模型参数。",
        width=6.25,
    )


def part3(doc, data):
    h1(doc, "第三部分：实验数据与预处理")
    overview = data["dataset"]
    rows = [
        ("数据集名称", "GTSRB German Traffic Sign Recognition Benchmark"),
        ("类别数", 43),
        ("训练集图片数", overview["train_images"]),
        ("测试集图片数", overview["test_images"]),
        ("总图片数", overview["total_images"]),
        ("图像尺寸范围", f"宽 {overview['min_width']}-{overview['max_width']}，高 {overview['min_height']}-{overview['max_height']}"),
    ]
    table(
        doc,
        "GTSRB 数据集基本信息",
        ["项目", "数值"],
        rows,
        [2600, 6760],
        "本表对应老师模板中的“数据集选择”。GTSRB 是公开交通标志识别数据集，类别覆盖限速、禁止、警告、指示等标志，适合作为交通标志分类和鲁棒性实验基础。",
    )
    figure(
        doc,
        "fig_01_class_distribution.png",
        "GTSRB 类别分布统计",
        "该图展示每个类别的样本数量。类别不平衡会影响模型训练：样本多的类别更容易学到稳定特征，样本少的类别更容易误判。因此报告后续使用 macro precision/recall/F1，而不仅只看总体 accuracy。",
        width=6.05,
    )
    figure(
        doc,
        "fig_02_image_size_distribution.png",
        "GTSRB 图像尺寸分布",
        "该图说明原始图像尺寸并不统一。为了让模型 batch 输入一致，本项目统一 resize 到 64 x 64。选择 64 x 64 是精度和计算量的折中：尺寸太小会丢失标志细节，尺寸太大则训练和攻击评估更慢。",
        width=5.95,
    )
    figure(
        doc,
        "fig_03_train_samples.png",
        "训练集样例图",
        "该图用于展示数据集真实外观，包括不同光照、模糊、角度和背景。可以看到交通标志常常较小且存在拍摄质量差异，这也是需要数据增强和鲁棒性分析的原因。",
        width=6.20,
    )
    figure(
        doc,
        "fig_04_test_samples.png",
        "测试集样例图",
        "该图展示测试样本，与训练样本类似但不参与模型参数更新。所有 clean、attack、defense 指标都基于测试集或测试集分层抽样，保证评估对象独立。",
        width=6.20,
    )

    rows = [
        ("Resize", "统一到 64 x 64", "保证输入尺寸一致"),
        ("Random Rotation", "随机旋转 10 度以内", "模拟拍摄角度变化"),
        ("Color Jitter", "亮度、对比度、饱和度扰动", "模拟光照变化"),
        ("Random Affine", "小幅平移和缩放", "模拟位置和尺度变化"),
        ("Normalize", "mean=(0.3337,0.3064,0.3171)，std=(0.2672,0.2564,0.2629)", "让输入分布更稳定；攻击 epsilon 在此空间定义"),
        ("Selected indices", "固定 3000 张分层抽样测试图", "让攻击、防御和补充实验使用同一批样本，结果可比"),
    ]
    table(
        doc,
        "数据预处理与增强策略",
        ["步骤/参数", "设置", "为什么这样做"],
        rows,
        [2300, 3500, 3560],
        "本表解释每个预处理参数。数据增强只用于训练阶段，测试、攻击和防御评估阶段不使用随机增强，以免评估结果不稳定。",
    )


def part4(doc, data):
    h1(doc, "第四部分：实验结果与性能评估")
    h2(doc, "4.1 基础识别结果")
    baseline = data["baseline"]
    rows = [
        ("Accuracy", pct(baseline["accuracy"]), "总体分类正确率"),
        ("Precision Macro", num(baseline["precision_macro"]), "每类 precision 平均"),
        ("Recall Macro", num(baseline["recall_macro"]), "每类 recall 平均"),
        ("F1 Macro", num(baseline["f1_macro"]), "precision 与 recall 综合指标"),
        ("Best Epoch", baseline["best_epoch"], "验证集表现最佳的训练轮次"),
        ("Parameter Count", f"{baseline['parameter_count']:,}", "可训练参数量"),
    ]
    table(
        doc,
        "ResNet18 基础识别测试结果",
        ["指标", "数值", "解释"],
        rows,
        [2600, 2300, 4460],
        "本表是 clean baseline 的核心结果。Accuracy 达到 97.81%，说明基础识别模型已经足够稳定，后续对抗攻击导致的下降更能说明鲁棒性问题，而不是模型本身没有学会任务。",
    )
    figure(
        doc,
        "fig_05_baseline_loss_curve.png",
        "ResNet18 训练损失曲线",
        "该图展示训练集和验证集 loss 随 epoch 变化。loss 下降表示模型逐渐学会分类边界；如果训练 loss 降而验证 loss 升，则可能过拟合。本项目曲线整体收敛，说明模型训练过程有效。",
        width=5.70,
    )
    figure(
        doc,
        "fig_06_baseline_accuracy_curve.png",
        "ResNet18 训练与验证准确率曲线",
        "该图用于观察模型泛化能力。训练准确率和验证准确率同步提升，说明数据增强与 ResNet18 结构能够有效学习交通标志特征。最佳 epoch 用于保存 best_model.pth，后续攻击和防御都使用该 checkpoint。",
        width=5.70,
    )
    figure(
        doc,
        "fig_07_baseline_confusion_matrix.png",
        "测试集混淆矩阵",
        "混淆矩阵的横纵轴分别表示预测类别和真实类别，对角线越亮表示正确分类越多。非对角线亮点代表误判来源，可用于发现哪些类别容易混淆，例如形状相近、数字相近或图像质量较差的标志。",
        width=5.90,
    )
    figure(
        doc,
        "fig_08_baseline_per_class_accuracy.png",
        "各类别测试准确率",
        "该图按类别展示准确率，比总体 accuracy 更细。某些类别样本少或视觉相似度高，可能低于整体水平；这说明评价模型时不能只看一个平均数，还要看类别级表现。",
        width=6.05,
    )
    figure(
        doc,
        "fig_09_baseline_correct_samples.png",
        "基础模型正确识别样例",
        "该图展示模型在正常样本上的成功案例，用于说明模型学到的特征是有效的，例如圆形限速牌、三角警告牌和蓝色指示牌都能被正确分类。",
        width=6.20,
    )
    figure(
        doc,
        "fig_10_baseline_wrong_samples.png",
        "基础模型错误识别样例",
        "该图是失败案例分析。错误通常出现在低光照、模糊、遮挡、类别相似或标志较小的样本中。这些失败案例说明模型决策边界仍然存在脆弱区域，也为对抗攻击实验提供背景。",
        width=6.20,
    )

    h2(doc, "4.2 对抗攻击结果与指标分析")
    attack_df = data["attack"]
    rows = []
    for _, row in attack_df.sort_values(["attack", "epsilon"]).iterrows():
        rows.append([row.attack.upper(), num(row.epsilon, 3), pct(row.clean_accuracy), pct(row.adversarial_accuracy), pct(row.attack_success_rate), num(row.mean_confidence_drop, 4)])
    table(
        doc,
        "FGSM 与 PGD 攻击完整指标",
        ["Attack", "Epsilon", "Clean Acc.", "Adv. Acc.", "Success", "Conf. Drop"],
        rows,
        [1000, 1200, 1600, 1600, 1600, 2360],
        "本表比较不同 epsilon 下的攻击效果。Adv. Acc. 表示对抗样本准确率，Success 表示原本预测正确但攻击后预测错误的比例，Conf. Drop 表示置信度平均下降。随着 epsilon 增大，对抗准确率下降，说明扰动预算越大攻击越强。",
        font_size=8.4,
    )
    figure(
        doc,
        "fig_11_attack_accuracy_curve.png",
        "不同 epsilon 下对抗准确率变化",
        "该图直接展示攻击强度。FGSM 和 PGD 的准确率都随 epsilon 增大而下降；PGD 在多数 epsilon 下更低，因为它进行多步迭代优化，而不是单步更新。",
        width=5.95,
    )
    figure(
        doc,
        "fig_12_attack_success_curve.png",
        "不同 epsilon 下攻击成功率变化",
        "该图与准确率曲线互补。攻击成功率越高，说明越多原本识别正确的样本被扰动推成错误类别。PGD 成功率高于 FGSM，体现了多步投影梯度搜索的优势。",
        width=5.95,
    )
    figure(
        doc,
        "fig_13_attack_confidence_drop.png",
        "不同 epsilon 下平均置信度下降",
        "该图分析的不只是是否预测错误，还包括模型信心如何变化。置信度下降说明扰动让模型的决策变得不稳定，即使某些样本还没被完全攻破，也已经靠近决策边界。",
        width=5.95,
    )
    h2(doc, "4.2.1 两种攻击方式的结果解读")
    fgsm003 = attack_df[(attack_df.attack == "fgsm") & (np.isclose(attack_df.epsilon, 0.03))].iloc[0]
    pgd003 = attack_df[(attack_df.attack == "pgd") & (np.isclose(attack_df.epsilon, 0.03))].iloc[0]
    fgsm005 = attack_df[(attack_df.attack == "fgsm") & (np.isclose(attack_df.epsilon, 0.05))].iloc[0]
    pgd005 = attack_df[(attack_df.attack == "pgd") & (np.isclose(attack_df.epsilon, 0.05))].iloc[0]
    rows = [
        ("epsilon=0.03", "FGSM", pct(fgsm003.adversarial_accuracy), pct(fgsm003.attack_success_rate), "单步攻击已能显著降低准确率"),
        ("epsilon=0.03", "PGD", pct(pgd003.adversarial_accuracy), pct(pgd003.attack_success_rate), "PGD 更低，说明多步搜索找到更强扰动"),
        ("epsilon=0.05", "FGSM", pct(fgsm005.adversarial_accuracy), pct(fgsm005.attack_success_rate), "扰动预算增大后单步攻击也变强"),
        ("epsilon=0.05", "PGD", pct(pgd005.adversarial_accuracy), pct(pgd005.attack_success_rate), "PGD 仍保持强攻击，但与 FGSM 差距缩小"),
    ]
    table(
        doc,
        "FGSM 与 PGD 关键结果对照",
        ["场景", "攻击", "Adv. Acc.", "Success", "解读"],
        rows,
        [1500, 1000, 1600, 1600, 3660],
        "本表直接比较两种攻击。epsilon=0.03 时 PGD 准确率低于 FGSM，说明在相同最大扰动幅度下，多步投影优化更容易找到模型敏感方向。epsilon=0.05 时 FGSM 也很强，是因为扰动预算更大，单步方向已经足以跨过更多决策边界。",
        font_size=8.2,
    )
    add_para(
        doc,
        "从算法角度解释：FGSM 只在原始点 x 处看一次梯度，等价于用局部线性近似估计最坏方向；如果模型损失曲面在 epsilon 邻域内并非严格线性，单步近似就可能不够充分。PGD 每次更新后都在新点重新计算梯度，因此能沿着损失曲面逐步逼近更高损失区域。这也是 PGD 在鲁棒性评估中更常作为强攻击基线的原因。",
    )

    h2(doc, "4.3 扰动可感知性：证明不是明显改图")
    pert_df = data["perturbation"]
    rows = []
    for _, row in pert_df.sort_values(["attack", "epsilon_normalized_space"]).iterrows():
        rows.append([row.attack.upper(), num(row.epsilon_normalized_space, 3), num(row.mean_linf_pixel_space, 6), num(row.mean_abs_delta_pixel_space, 6), f"{row.mean_psnr_db:.2f}", f"{int(row.successful_changes)}/{int(row.total_images)}"])
    table(
        doc,
        "像素空间扰动可感知性指标",
        ["Attack", "Norm eps", "Pixel Linf", "Mean |Delta|", "PSNR", "Changed"],
        rows,
        [1000, 1200, 1800, 1800, 1500, 2060],
        "本表把归一化空间 epsilon 转换到像素空间。Pixel Linf 约 0.008 时，相当于约 2/255 的最大像素变化；PSNR 在 42 dB 以上，说明真实扰动较小。Changed 表示成功改变预测的样本数量。",
        font_size=8.4,
    )
    figure(
        doc,
        "fig_20_perturbation_psnr_bar.png",
        "不同攻击设置下平均 PSNR",
        "PSNR 越高表示扰动越不明显。该图显示 epsilon=0.01 的 PSNR 约 52 dB，epsilon=0.03 约 42-44 dB。即使 epsilon=0.03，扰动仍然属于小幅像素变化，而不是肉眼明显涂改。",
        width=5.70,
    )
    figure(
        doc,
        "fig_21_fgsm_perceptual_grid.png",
        "FGSM epsilon=0.03 真实对抗图与放大扰动视图",
        "每组图包含 Original、Adversarial、Delta x30 和 Delta heat。Original 与 Adversarial 的视觉差异很小；Delta x30 和 Delta heat 是人为放大 30 倍或归一化后的分析图，用于看扰动位置，不代表真实扰动幅度。",
        width=6.25,
    )
    figure(
        doc,
        "fig_22_pgd_perceptual_grid.png",
        "PGD epsilon=0.03 真实对抗图与放大扰动视图",
        "该图与 FGSM 图对照。PGD 的真实对抗图同样与原图接近，但预测已经改变，说明攻击利用模型梯度敏感方向，而不是制造肉眼可见的大噪声。",
        width=6.25,
    )

    h2(doc, "4.4 消融实验一：随机噪声对照")
    random_df = data["random_noise"]
    fgsm = attack_df[attack_df.attack == "fgsm"].set_index("epsilon")
    pgd = attack_df[attack_df.attack == "pgd"].set_index("epsilon")
    rows = []
    for _, row in random_df.sort_values("epsilon").iterrows():
        eps = float(row.epsilon)
        rows.append([num(eps, 3), pct(row.perturbed_accuracy), pct(fgsm.loc[eps].adversarial_accuracy), pct(pgd.loc[eps].adversarial_accuracy), pct(row.success_rate_on_clean_correct)])
    table(
        doc,
        "随机 Linf 噪声与梯度攻击对照",
        ["Epsilon", "Random Acc.", "FGSM Acc.", "PGD Acc.", "Random Success"],
        rows,
        [1400, 1900, 1900, 1900, 2260],
        "本表是证明攻击有效性的关键对照。随机噪声与 FGSM/PGD 使用相同 epsilon 预算，但随机扰动几乎不降低准确率，而梯度攻击显著降低准确率。这说明对抗攻击不是普通噪声，而是利用梯度方向的定向扰动。",
        font_size=8.4,
    )
    figure(
        doc,
        "fig_27_random_vs_adversarial_accuracy.png",
        "随机噪声对照与梯度攻击准确率对比",
        "蓝线代表随机扰动，几乎保持在 97.9% 附近；橙线和绿线代表 FGSM/PGD，随 epsilon 增大明显下降。这张图直接回应“攻击是不是太明显或只是加噪声”的问题：同等幅度随机噪声无效，梯度方向才有效。",
        width=6.05,
    )
    figure(
        doc,
        "fig_32_random_vs_pgd_visual_grid.png",
        "随机噪声与 PGD 对抗样本视觉对照",
        "该图在同一组样本上比较 Original、Random、PGD 和 PGD delta x30。Random 图像和 PGD 图像都与原图接近，但 PGD 更容易改变预测，说明视觉相似不等于模型决策稳定。",
        width=6.25,
    )

    h2(doc, "4.5 消融实验二：PGD 迭代步数")
    step_df = data["pgd_steps"]
    rows = []
    for _, row in step_df.sort_values("pgd_steps").iterrows():
        rows.append([int(row.pgd_steps), pct(row.adversarial_accuracy), pct(row.attack_success_rate), num(row.mean_confidence_drop, 4)])
    table(
        doc,
        "PGD 迭代步数消融，epsilon=0.03",
        ["Steps", "Adv. Acc.", "Attack Success", "Conf. Drop"],
        rows,
        [1600, 2200, 2500, 3060],
        "本表固定 epsilon=0.03 和 alpha=0.005，只改变 PGD steps。steps 是迭代次数，表示在同一个扰动球内搜索多少次。步数增加后对抗准确率下降，说明 PGD 强度来自更充分的优化，而不是更大的扰动幅度。",
        font_size=8.4,
    )
    figure(
        doc,
        "fig_28_pgd_step_ablation.png",
        "PGD 步数对攻击强度的影响",
        "图中蓝线是对抗准确率，红线是攻击成功率。随着 steps 从 1 增加到 20，模型准确率显著下降，攻击成功率上升。这解释了为什么 PGD 通常被视为比 FGSM 更强的白盒攻击基线。",
        width=5.90,
    )

    h2(doc, "4.6 决策边界分析：Top-1/Top-2 概率间隔")
    margin_df = data["margin_shift"]
    rows = []
    for _, row in margin_df.iterrows():
        rows.append([row.condition, num(row.mean_top1_top2_margin, 4), num(row.median_top1_top2_margin, 4), pct(row["low_margin_rate_lt_0.20"])])
    table(
        doc,
        "概率间隔变化",
        ["Condition", "Mean Margin", "Median Margin", "Low-margin Rate"],
        rows,
        [2600, 2200, 2200, 2360],
        "Margin 是 top-1 概率减去 top-2 概率。margin 越小，说明模型第一选择和第二选择越接近，决策越不稳定。FGSM/PGD 会压缩 margin，而随机扰动几乎不改变 margin。",
        font_size=8.4,
    )
    figure(
        doc,
        "fig_31_margin_shift_histogram.png",
        "Clean、随机扰动和对抗扰动下的概率间隔分布",
        "该图从概率分布角度解释攻击效果。Clean 和 Random 曲线接近，说明随机扰动没有显著推动样本靠近决策边界；FGSM/PGD 曲线向低 margin 区域移动，说明对抗扰动会削弱模型决策稳定性。",
        width=5.95,
    )

    h2(doc, "4.7 输入防御结果与参数扫描")
    defense_df = data["defense"]
    rows = []
    for attack in ["fgsm", "pgd"]:
        for eps in [0.03, 0.05]:
            subset = defense_df[(defense_df.attack == attack) & (np.isclose(defense_df.epsilon, eps))]
            before = subset.iloc[0].adversarial_accuracy_before_defense
            for _, row in subset.sort_values("defense").iterrows():
                rows.append([attack.upper(), num(eps, 2), row.defense, pct(before), pct(row.defended_accuracy), pct(row.recovery_rate_on_successful_attacks)])
    table(
        doc,
        "固定参数输入防御结果",
        ["Attack", "Eps", "Defense", "Before", "After", "Recovery"],
        rows,
        [1000, 900, 2300, 1500, 1500, 2160],
        "本表比较 Gaussian Blur、Median Filter 和 JPEG Compression 三种防御。Before 是防御前对抗准确率，After 是防御后准确率，Recovery 是对成功攻击样本恢复正确的比例。JPEG 在多数设置下最好。",
        font_size=8.2,
    )
    figure(
        doc,
        "fig_16_fgsm_input_defense_curve.png",
        "FGSM 攻击下输入防御效果",
        "该图展示 FGSM eps=0.03 和 0.05 下三种防御的准确率变化。防御曲线高于 before_defense 说明输入变换确实恢复了部分样本，但仍未恢复到 clean accuracy。",
        width=5.95,
    )
    figure(
        doc,
        "fig_17_pgd_input_defense_curve.png",
        "PGD 攻击下输入防御效果",
        "PGD 比 FGSM 更强，因此防御恢复更困难。图中 JPEG Compression 依旧表现较好，说明压缩高频细节对当前扰动有明显削弱作用。",
        width=5.95,
    )
    figure(
        doc,
        "fig_18_input_defense_accuracy_bar.png",
        "epsilon=0.03 下不同输入防御横向对比",
        "该柱状图把同一 epsilon 下的防御方法放在一起比较，便于展示哪种方法恢复最多。横向对比比单独曲线更适合答辩时快速说明结论。",
        width=6.20,
    )
    figure(
        doc,
        "fig_19_input_defense_examples.png",
        "输入防御恢复样例",
        "该图展示 Original、Attack 和 Defense 三列。它说明防御不是重新训练模型，而是对同一对抗图像做输入变换后重新识别；绿色结果表示防御恢复了正确预测。",
        width=6.20,
    )
    h2(doc, "4.7.1 防御效果的原理解读")
    fixed_rows = []
    for attack_name, eps in [("fgsm", 0.03), ("pgd", 0.03), ("pgd", 0.05)]:
        subset = defense_df[(defense_df.attack == attack_name) & (np.isclose(defense_df.epsilon, eps))]
        before = subset.iloc[0].adversarial_accuracy_before_defense
        best = subset.sort_values("defended_accuracy", ascending=False).iloc[0]
        fixed_rows.append(
            [
                attack_name.upper(),
                num(eps, 2),
                pct(before),
                best.defense,
                pct(best.defended_accuracy),
                pct(best.defended_accuracy - before),
            ]
        )
    table(
        doc,
        "防御恢复效果摘要",
        ["Attack", "Eps", "Before", "Best Defense", "After", "Gain"],
        fixed_rows,
        [1000, 900, 1500, 2500, 1500, 1960],
        "本表把防御结果压缩成最关键的恢复量。JPEG Compression 在多个场景下表现最好，说明当前对抗扰动中有相当一部分可被有损压缩破坏。但 Gain 仍不能让模型完全回到 clean accuracy，说明输入预处理只能部分防御。",
        font_size=8.2,
    )
    add_para(
        doc,
        "为什么 JPEG Compression 效果较好：JPEG 会把图像分块后做 DCT 频域变换，再对高频系数量化。许多对抗扰动虽然在像素幅度上很小，但会表现为高频、局部、方向性的变化；压缩过程会丢弃一部分高频细节，因此能破坏扰动结构。为什么 Gaussian Blur 和 Median Filter 也有效但较弱：它们主要在空间域平滑或替换局部像素，对某些高频扰动有抑制作用，但可能无法完全消除梯度攻击形成的全局方向性变化。",
    )
    add_para(
        doc,
        "防御局限也必须说明：输入预处理是在模型外部改变输入，不改变 ResNet18 的决策边界。如果攻击者知道防御存在，可以设计自适应攻击，把 JPEG 或滤波近似纳入攻击流程，防御效果可能下降。因此本项目把输入预处理作为轻量防御基线，而不是最终安全方案。",
    )

    sweep_df = data["defense_sweep"].sort_values("pgd_defended_accuracy", ascending=False)
    rows = []
    for _, row in sweep_df.iterrows():
        rows.append([row.variant, pct(row.clean_transformed_accuracy), pct(row.pgd_defended_accuracy), pct(row.clean_accuracy_drop), pct(row.pgd_accuracy_gain)])
    table(
        doc,
        "输入防御参数扫描",
        ["Variant", "Clean Acc.", "PGD Def. Acc.", "Clean Drop", "PGD Gain"],
        rows,
        [2500, 1700, 1900, 1500, 1760],
        "本表是防御消融实验。Clean Drop 表示防御对正常样本准确率造成的损失，PGD Gain 表示防御对 PGD 样本准确率的提升。jpeg_q50 鲁棒性最高，但 clean accuracy 损失也更大；jpeg_q75 更均衡。",
        font_size=8.2,
    )
    figure(
        doc,
        "fig_29_defense_parameter_tradeoff.png",
        "防御参数扫描：clean accuracy 与 robust accuracy 权衡",
        "横轴是正常样本经过预处理后的准确率，纵轴是 PGD eps=0.03 防御后准确率。理想点在右上角。图中可见更强压缩或更大滤波窗口通常提高鲁棒性，但会向左移动，表示 clean accuracy 降低。",
        width=6.05,
    )
    figure(
        doc,
        "fig_30_defense_parameter_bar.png",
        "PGD epsilon=0.03 下不同防御参数准确率对比",
        "该图按防御后准确率排序，展示各参数的鲁棒性能。虚线是防御前 PGD 准确率，柱子越高表示恢复越多。它直观说明参数选择会显著影响防御效果。",
        width=6.05,
    )

    h2(doc, "4.7.2 JPEG Quality 消融实验：防御强度与正常识别代价")
    jpeg_df = data["jpeg_quality"]
    rows = []
    for attack_name, eps in [("clean", 0.0), ("fgsm", 0.03), ("pgd", 0.03), ("fgsm", 0.05), ("pgd", 0.05)]:
        subset = jpeg_df[(jpeg_df.attack == attack_name) & (np.isclose(jpeg_df.epsilon, eps))]
        for _, row in subset.sort_values("jpeg_quality").iterrows():
            rows.append(
                [
                    attack_name.upper(),
                    num(eps, 2),
                    f"Q{int(row.jpeg_quality)}",
                    pct(row.accuracy_before_jpeg),
                    pct(row.accuracy_after_jpeg),
                    pct(row.recovery_rate_on_successful_attacks),
                    num(row.mean_psnr_vs_attack, 2),
                ]
            )
    table(
        doc,
        "JPEG Quality 防御消融结果",
        ["Attack", "Eps", "Quality", "Before", "After", "Recovery", "PSNR"],
        rows,
        [900, 800, 1050, 1350, 1350, 1600, 2310],
        "本表单独扫描 JPEG quality=50、75、90。Quality 越低，压缩越强，高频扰动被破坏得更多，但正常图像细节也更容易损失。clean 行用于观察正常识别代价；FGSM/PGD 行用于观察鲁棒恢复。结果显示 Q50 鲁棒恢复最强，Q75 在 clean accuracy 与 defended accuracy 之间更均衡。",
        font_size=7.6,
    )
    figure(
        doc,
        "fig_34_fgsm_jpeg_quality_curve.png",
        "FGSM 下 JPEG Quality 对防御效果的影响",
        "该图展示 FGSM epsilon=0.03 与 0.05 下，不同 JPEG quality 的防御后准确率。Q50 曲线最高，说明强压缩对 FGSM 扰动破坏更明显；但这并不意味着 Q50 一定最适合部署，还需要结合 clean accuracy 代价判断。",
        width=5.95,
    )
    figure(
        doc,
        "fig_35_pgd_jpeg_quality_curve.png",
        "PGD 下 JPEG Quality 对防御效果的影响",
        "该图展示 PGD 攻击下的 JPEG quality 消融。PGD 比 FGSM 更强，但 Q50/Q75 仍能显著高于 before JPEG，说明 JPEG 压缩确实破坏了部分对抗扰动结构。Q90 保留细节更多，clean 损失小，但鲁棒恢复也较弱。",
        width=5.95,
    )
    figure(
        doc,
        "fig_36_jpeg_quality_eps003_bar.png",
        "epsilon=0.03 下 JPEG Quality 横向对比",
        "该柱状图把 FGSM 与 PGD 在 epsilon=0.03 下的 Q50/Q75/Q90 放在一起比较。它清楚展示 Q50 防御准确率最高、Q75 次之、Q90 最低，说明防御强度与压缩强度高度相关。",
        width=5.95,
    )
    figure(
        doc,
        "fig_37_jpeg_clean_accuracy_cost.png",
        "JPEG Quality 对正常样本准确率的影响",
        "该图专门展示 clean 输入经过 JPEG 后的准确率。Q90 最接近原始 clean accuracy，Q50 损失最大。将本图与防御准确率曲线一起看，可以说明防御不是越强越好，而是需要在正常识别与鲁棒恢复之间取折中。",
        width=5.75,
    )
    figure(
        doc,
        "fig_38_fgsm_jpeg_quality_examples.png",
        "FGSM 攻击下 JPEG Quality 防御案例",
        "该图按 Clean、Attack、JPEG Q50、Q75、Q90 展示同一批样本。它能直观看到：强压缩会改变图像纹理细节，但也更可能把模型预测从攻击错误类别拉回正确类别。该图适合答辩时解释参数消融不是只看数字，也有可视化证据。",
        width=6.20,
    )
    figure(
        doc,
        "fig_39_pgd_jpeg_quality_examples.png",
        "PGD 攻击下 JPEG Quality 防御案例",
        "该图展示更强 PGD 攻击下不同 JPEG quality 的恢复样例。PGD 扰动更贴近模型局部最坏方向，因此恢复更困难；若防御后仍能恢复，说明输入压缩确实改变了模型敏感的局部高频结构。",
        width=6.20,
    )

    h2(doc, "4.7.3 Grad-CAM 可解释性分析：攻击如何改变模型关注区域")
    gradcam_df = data["gradcam"]
    rows = []
    for _, row in gradcam_df.iterrows():
        rows.append(
            [
                int(row.case_id),
                f"{int(row.label)} {row.label_name}",
                int(row.clean_pred),
                int(row.adv_pred),
                int(row.def_pred),
                num(row.clean_conf, 3),
                num(row.adv_conf, 3),
                num(row.def_conf, 3),
            ]
        )
    table(
        doc,
        "Grad-CAM 案例预测记录",
        ["Case", "True Label", "Clean", "Attack", "Defense", "C Conf.", "A Conf.", "D Conf."],
        rows,
        [700, 2500, 800, 850, 900, 1100, 1100, 1410],
        "本表记录 Grad-CAM 图中每个案例的预测变化。6 个案例都满足 clean 识别正确、PGD 攻击后预测错误、JPEG 防御后恢复正确。Conf. 是模型对当前预测类别的置信度，用于说明攻击并非只造成轻微不确定，而是能让模型较高置信地转向错误类别。",
        font_size=7.5,
    )
    figure(
        doc,
        "fig_33_gradcam_clean_attack_defense.png",
        "Clean、Attack、Defense 三阶段 Grad-CAM 对比",
        "Grad-CAM 用最后一层卷积特征的梯度权重生成热力图，红色区域表示模型更依赖的图像区域。图中 Attack 列的关注区域相比 Clean 列出现扩散、偏移或关注错误纹理；Defense 列在 JPEG 处理后预测恢复，热力图也更接近标志主体区域。该图补充了纯指标无法说明的模型内部证据。",
        width=5.95,
    )

    h2(doc, "4.7.4 全测试集验证与运行效率评估")
    full_attack_df = data["full_attack"]
    rows = []
    for _, row in full_attack_df.iterrows():
        rows.append(
            [
                row.attack.upper(),
                num(row.epsilon, 2),
                int(row.total_images),
                pct(row.clean_accuracy),
                pct(row.adversarial_accuracy),
                pct(row.attack_success_rate),
                num(row.mean_confidence_drop, 4),
            ]
        )
    table(
        doc,
        "全测试集攻击验证结果",
        ["Attack", "Eps", "Images", "Clean Acc.", "Adv. Acc.", "Success", "Conf. Drop"],
        rows,
        [900, 800, 1100, 1500, 1500, 1500, 2060],
        "主实验使用固定 3000 张分层抽样测试图以控制运行时间。该表进一步在完整 GTSRB 测试集 12630 张图上验证 epsilon=0.03 的关键结论。FGSM 与 PGD 的全测试集结果和抽样趋势一致，说明前文结论不是抽样偶然。",
        font_size=8.0,
    )
    full_def_df = data["full_defense"]
    rows = []
    for _, row in full_def_df.iterrows():
        rows.append(
            [
                row.attack.upper(),
                num(row.epsilon, 2),
                row.defense,
                int(row.total_images),
                pct(row.adversarial_accuracy_before_defense),
                pct(row.defended_accuracy),
                pct(row.recovery_rate_on_successful_attacks),
            ]
        )
    table(
        doc,
        "全测试集 JPEG Q75 防御验证结果",
        ["Attack", "Eps", "Defense", "Images", "Before", "After", "Recovery"],
        rows,
        [900, 800, 2100, 1100, 1350, 1350, 1760],
        "本表在完整测试集上只验证主防御配置 JPEG Q75。FGSM epsilon=0.03 从 66.67% 恢复到 78.35%，PGD epsilon=0.03 从 59.25% 恢复到 78.04%。这说明 JPEG 防御在完整测试集上仍能稳定恢复一部分对抗样本。",
        font_size=8.0,
    )
    runtime_df = data["runtime"].sort_values("mean_ms_per_image")
    rows = []
    for _, row in runtime_df.iterrows():
        rows.append([row.operation, int(row.batches), num(row.mean_ms_per_image, 3), num(row.mean_images_per_second, 1)])
    table(
        doc,
        "推理、攻击与防御运行效率统计",
        ["Operation", "Batches", "ms/image", "Images/s"],
        rows,
        [3300, 1200, 1800, 3060],
        "本表说明不同模块的计算代价。clean inference 与 JPEG defense 都较快；FGSM 需要一次反向传播，耗时增加；PGD 需要多步反向传播，因此最慢。该结果可用于解释为什么 PGD 更强但成本更高，也说明 JPEG 输入防御适合做轻量展示模块。",
        font_size=8.0,
    )
    figure(
        doc,
        "fig_40_runtime_ms_per_image.png",
        "不同操作的单张图片平均耗时",
        "该图用横向柱状图展示 ms/image。PGD generation 明显高于 FGSM generation 和 clean inference，原因是 PGD 要进行多步梯度更新。JPEG defense plus inference 的耗时较低，说明输入预处理防御不会带来很大的演示负担。",
        width=5.95,
    )
    figure(
        doc,
        "fig_41_runtime_images_per_second.png",
        "不同操作的处理吞吐率",
        "吞吐率与单图耗时互为补充。clean inference 和 JPEG defense 的 images/s 较高，而 PGD generation 最低。该图适合在 Demo 部分说明：实时演示可以优先展示 clean、FGSM 和 JPEG 防御，PGD 可作为预生成或短样本演示。",
        width=5.95,
    )

    h2(doc, "4.8 深度反思：成功、失败与原因")
    rows = [
        ("正常识别成功", "边界清晰、光照正常、类别特征明显", "ResNet 能稳定提取形状和颜色特征"),
        ("正常识别失败", "低光照、模糊、遮挡、类别相似", "图像信息不足或决策边界接近"),
        ("攻击成功", "视觉变化小但预测改变", "梯度扰动沿损失增大方向移动样本"),
        ("随机噪声失败", "同等 epsilon 随机扰动几乎不降准确率", "随机方向大多不是模型敏感方向"),
        ("防御恢复", "JPEG/滤波削弱高频扰动", "输入变换使样本回到较稳定区域"),
        ("防御失败", "强攻击或过强压缩导致细节丢失", "轻量输入防御有上限，自适应攻击可能绕过"),
    ]
    table(
        doc,
        "典型现象与原因分析",
        ["现象", "表现", "原因"],
        rows,
        [1900, 3300, 4160],
        "本表对应老师模板中的“深度反思”。报告不能只展示好结果，也要说明失败样例和方法局限。当前防御是轻量基线，不等于完整安全方案。",
    )


def part5(doc):
    h1(doc, "第五部分：Demo 演示设计")
    add_para(
        doc,
        "老师模板中 YOLO 项目可以播放实时检测录屏。本项目是交通标志分类与鲁棒性分析，适合演示“上传或选择交通标志图片 -> ResNet clean 识别 -> 生成 FGSM/PGD 对抗样本 -> 展示扰动放大图 -> 输入防御恢复”的流程。当前仓库已经保存了可直接展示的对抗样例、扰动热力图和防御恢复图，具备离线演示材料。",
    )
    rows = [
        ("步骤 1", "展示 clean 识别", "使用 fig_09 正确样例和 fig_10 错误样例说明基础模型能力"),
        ("步骤 2", "展示攻击生成", "展示 fig_21/fig_22，强调真实对抗图和放大 Delta 的区别"),
        ("步骤 3", "展示随机噪声对照", "展示 fig_27/fig_32，证明不是普通噪声"),
        ("步骤 4", "展示防御恢复", "展示 fig_19 和 fig_30，说明输入预处理恢复效果"),
        ("步骤 5", "展示 Grad-CAM", "展示 fig_33，说明攻击前后模型关注区域如何变化"),
        ("步骤 6", "讲解运行代价", "展示 fig_40/fig_41，说明 PGD 更强但更慢，JPEG 防御开销较低"),
        ("步骤 7", "讲解局限", "说明输入防御不是完备方案，后续可做对抗训练和自适应攻击"),
    ]
    table(
        doc,
        "Demo 演示脚本",
        ["演示步骤", "内容", "讲解重点"],
        rows,
        [1500, 2900, 4960],
        "本表给出答辩 Demo 流程。当前没有生成实时录屏，因此报告不虚构实时视频；但已有图片和脚本可以支撑一次完整的离线演示。",
    )
    h2(doc, "可复现命令")
    for cmd in [
        "python -m src.train_classifier --config configs/baseline_resnet18.yaml",
        "python -m src.evaluate_attacks --config configs/attack_fgsm_pgd.yaml",
        "python -m src.analyze_perturbations",
        "python -m src.evaluate_input_defense --config configs/defense_input_preprocessing.yaml",
        "python -m src.extended_experiments",
        "python -m src.visualization.gradcam_analysis --config configs/explainability_gradcam.yaml",
        "python -m src.evaluate_jpeg_ablation --config configs/defense_jpeg_ablation.yaml",
        "python -m src.benchmark_runtime --config configs/runtime_benchmark.yaml",
        "python -m src.evaluate_attacks --config configs/attack_fgsm_pgd_full_test.yaml",
        "python -m src.evaluate_input_defense --config configs/defense_input_preprocessing_full_test.yaml",
        "python scripts/build_teacher_format_report.py",
    ]:
        add_code_line(doc, cmd)
    explain(
        doc,
        "命令解释",
        "这些命令从训练、攻击、扰动分析、防御评估、补充消融到 Word 报告生成构成完整复现链路。若答辩时间有限，可不重新训练，只展示已保存的 results 和 reports 产物。",
    )


def part6(doc, data):
    h1(doc, "第六部分：总结与心得")
    attack_df = data["attack"]
    random_df = data["random_noise"]
    step_df = data["pgd_steps"]
    sweep_df = data["defense_sweep"]
    pgd003 = attack_df[(attack_df.attack == "pgd") & (np.isclose(attack_df.epsilon, 0.03))].iloc[0]
    random003 = random_df[np.isclose(random_df.epsilon, 0.03)].iloc[0]
    best_def = sweep_df.sort_values("pgd_defended_accuracy", ascending=False).iloc[0]
    full_pgd003 = data["full_attack"][(data["full_attack"].attack == "pgd") & (np.isclose(data["full_attack"].epsilon, 0.03))].iloc[0]
    q75_pgd003 = data["jpeg_quality"][
        (data["jpeg_quality"].attack == "pgd")
        & (np.isclose(data["jpeg_quality"].epsilon, 0.03))
        & (data["jpeg_quality"].jpeg_quality == 75)
    ].iloc[0]
    pgd_runtime = data["runtime"][data["runtime"].operation == "pgd_generation"].iloc[0]
    rows = [
        ("基础识别", "ResNet18 clean accuracy", pct(data["baseline"]["accuracy"]), "基础模型有效"),
        ("攻击有效性", "PGD eps=0.03 adversarial accuracy", pct(pgd003.adversarial_accuracy), "小扰动显著破坏识别"),
        ("随机对照", "Random eps=0.03 accuracy", pct(random003.perturbed_accuracy), "证明不是普通噪声"),
        ("PGD 消融", "20 steps adversarial accuracy", pct(step_df.sort_values("pgd_steps").iloc[-1].adversarial_accuracy), "迭代越充分攻击越强"),
        ("防御扫描", f"Best variant {best_def.variant}", pct(best_def.pgd_defended_accuracy), "存在 clean-robust trade-off"),
        ("JPEG 消融", "PGD eps=0.03 JPEG Q75 defended accuracy", pct(q75_pgd003.accuracy_after_jpeg), "主防御参数有独立消融支撑"),
        ("全测试集验证", "PGD eps=0.03 on 12630 images", pct(full_pgd003.adversarial_accuracy), "抽样结论在完整测试集成立"),
        ("运行效率", "PGD generation ms/image", num(pgd_runtime.mean_ms_per_image, 3), "强攻击计算代价最高"),
    ]
    table(
        doc,
        "项目核心结论",
        ["模块", "指标", "结果", "结论"],
        rows,
        [1600, 2700, 2100, 2960],
        "本表汇总最重要的答辩结论。项目从 clean recognition 扩展到 attack、defense 和 ablation，体现了对算法机制和实验验证的完整理解。",
    )
    add_callout(
        doc,
        "项目心得",
        "本项目最大的收获是：深度学习模型的高准确率不等于高安全性。对抗攻击展示了模型决策边界的脆弱性，随机噪声对照证明了梯度方向的重要性，防御参数扫描和 JPEG quality 消融说明鲁棒性提升往往伴随正常性能损失。Grad-CAM 则从关注区域角度补充了模型内部证据。后续如果继续推进，应重点补充对抗训练、自适应攻击和交互式 Demo。",
    )

    h2(doc, "参考文献")
    for ref in [
        "Stallkamp J., Schlipsing M., Salmen J., Igel C. The German Traffic Sign Recognition Benchmark: A multi-class classification competition. IJCNN, 2011.",
        "He K., Zhang X., Ren S., Sun J. Deep Residual Learning for Image Recognition. CVPR, 2016.",
        "Goodfellow I., Shlens J., Szegedy C. Explaining and Harnessing Adversarial Examples. ICLR, 2015.",
        "Madry A., Makelov A., Schmidt L., Tsipras D., Vladu A. Towards Deep Learning Models Resistant to Adversarial Attacks. ICLR, 2018.",
    ]:
        add_para(doc, ref, after=4)


def build_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generate_diagrams()
    data = load_data()
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    part1(doc, data)
    part2(doc, data)
    part3(doc, data)
    part4(doc, data)
    part5(doc)
    part6(doc, data)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_docx())
