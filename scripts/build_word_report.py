import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
FIG_DIR = REPORT_DIR / "figures"
TABLE_DIR = REPORT_DIR / "tables"
OUT_DIR = REPORT_DIR / "word"
OUT_PATH = OUT_DIR / "基于对抗攻击与输入防御的交通标志识别实验报告.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFC7D1", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths_dxa):
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run_font(run, size=9.5)
        if idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = str(text)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths_dxa)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=MUTED)


def add_figure(doc, filename, caption, width=6.15):
    path = FIG_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("基于对抗攻击与输入防御的交通标志识别实验报告 | 第 ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(footer)
    r = footer.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("模式识别与智能系统\n期末大作业实验报告")
    set_run_font(r, size=22, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("基于对抗攻击与输入防御的交通标志识别系统设计与实现")
    set_run_font(r, size=18, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("以 GTSRB 数据集为例的模型鲁棒性分析")
    set_run_font(r, size=12, color=MUTED)

    rows = [
        ("课程名称", "模式识别与智能系统"),
        ("项目方向", "交通标志识别、对抗样本攻击、鲁棒防御"),
        ("数据集", "GTSRB German Traffic Sign Recognition Benchmark"),
        ("完成阶段", "基础识别、FGSM/PGD 攻击、输入预处理防御"),
        ("小组成员", "待填写"),
        ("提交日期", "2026 年 6 月"),
    ]
    add_table(doc, ["项目", "内容"], rows, [1800, 7560])
    doc.add_page_break()


def add_abstract(doc):
    doc.add_heading("摘要", level=1)
    add_body_paragraph(
        doc,
        "交通标志识别是自动驾驶、辅助驾驶和智能交通系统中的重要视觉感知任务。深度学习模型虽然能够在标准测试集上取得较高准确率，但在面对对抗扰动时可能出现预测错误，从而带来潜在安全风险。本项目以 GTSRB 交通标志识别数据集为研究对象，首先构建基于 ResNet18 的基础识别模型，然后使用 FGSM 和 PGD 方法生成对抗样本，分析模型在不同扰动强度下的鲁棒性变化，最后使用 Gaussian Blur、Median Filter 和 JPEG Compression 三种输入预处理方法进行防御实验。实验结果表明，ResNet18 在正常测试集上达到 97.81% 的准确率，但在 PGD epsilon=0.03 攻击下准确率下降至 60.00%。输入预处理防御可以部分恢复模型性能，其中 JPEG Compression 在 PGD epsilon=0.03 下将准确率提升至 79.43%。本项目验证了交通标志识别模型在对抗扰动下的脆弱性，并初步探索了轻量级输入防御方法的有效性。",
    )
    p = doc.add_paragraph()
    r = p.add_run("关键词：")
    set_run_font(r, size=11, bold=True)
    r = p.add_run("交通标志识别；GTSRB；ResNet18；对抗样本；FGSM；PGD；输入预处理防御")
    set_run_font(r, size=11)

    doc.add_heading("报告结构", level=1)
    for item in [
        "实验背景与任务定义",
        "数据集与预处理",
        "基础交通标志识别模型",
        "对抗攻击方法与实验结果",
        "输入预处理防御方法与实验结果",
        "程序实现说明、问题分析、总结与分工",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_background(doc):
    doc.add_heading("1. 实验背景与任务定义", level=1)
    doc.add_heading("1.1 研究背景", level=2)
    add_body_paragraph(
        doc,
        "交通标志识别是智能交通系统中的基础任务之一。在自动驾驶场景中，车辆需要准确识别限速、禁止通行、转向、让行等交通标志，从而辅助决策系统完成速度控制、路径规划和风险规避。近年来，卷积神经网络在交通标志识别任务中取得了较高准确率，使得基于深度学习的识别方法成为主流方案。",
    )
    add_body_paragraph(
        doc,
        "然而，深度神经网络存在对抗脆弱性。所谓对抗样本，是指在原始输入上加入人眼难以察觉的细微扰动，却能导致模型产生错误预测的样本。对于交通标志识别任务而言，如果模型在轻微扰动、压缩噪声、光照变化或恶意干扰下识别错误，可能会影响智能交通系统的安全性。因此，本项目不只关注交通标志识别准确率，还进一步分析模型在对抗攻击下的鲁棒性。",
    )
    doc.add_heading("1.2 实验目标", level=2)
    for item in [
        "构建交通标志识别基础模型，并评估其正常识别性能。",
        "实现 FGSM 和 PGD 对抗攻击方法，生成对抗样本。",
        "分析不同扰动强度 epsilon 对模型准确率和攻击成功率的影响。",
        "使用输入预处理方法进行防御实验，比较不同防御策略的效果。",
        "保存完整实验指标、中间过程图片和可视化结果，为报告和答辩展示提供材料。",
    ]:
        add_numbered(doc, item)

    doc.add_heading("1.3 技术路线", level=2)
    add_body_paragraph(doc, "本项目整体流程为：GTSRB 数据集检查与可视化 -> ResNet18 基础识别模型训练 -> FGSM/PGD 对抗样本生成 -> 对抗攻击效果分析 -> Gaussian Blur、Median Filter、JPEG Compression 输入防御 -> 防御效果对比与总结。")


def add_dataset(doc):
    doc.add_heading("2. 数据集与预处理", level=1)
    doc.add_heading("2.1 GTSRB 数据集", level=2)
    add_body_paragraph(
        doc,
        "本项目使用 GTSRB 交通标志识别数据集。该数据集包含 43 类交通标志，图像来自真实道路场景，存在不同光照、角度、模糊和尺寸变化，适合用于交通标志分类实验。",
    )
    overview = json.load(open(TABLE_DIR / "table_02_dataset_overview.json", encoding="utf-8"))
    rows = [
        ("类别数", 43),
        ("训练集图片数", overview["train_images"]),
        ("测试集图片数", overview["test_images"]),
        ("总图片数", overview["total_images"]),
        ("图片宽度范围", f"{overview['min_width']} - {overview['max_width']}"),
        ("图片高度范围", f"{overview['min_height']} - {overview['max_height']}"),
    ]
    add_table(doc, ["项目", "数值"], rows, [2600, 6760])
    add_caption(doc, "表 1  GTSRB 数据集基本统计")
    add_figure(doc, "fig_01_class_distribution.png", "图 1  GTSRB 类别分布统计", width=6.25)
    add_figure(doc, "fig_03_train_samples.png", "图 2  GTSRB 训练集样例图", width=6.25)

    doc.add_heading("2.2 数据预处理", level=2)
    add_body_paragraph(
        doc,
        "实验中将输入图像统一缩放为 64 x 64。训练阶段使用随机旋转、颜色扰动和随机仿射等数据增强方法，以提高模型对视角和光照变化的适应能力。图像随后转换为 Tensor，并使用固定均值和标准差进行归一化。训练集进一步划分为训练子集和验证子集，其中验证比例为 15%。测试集用于最终性能评估和攻击防御实验。",
    )
    rows = [
        ("Resize", "统一缩放到 64 x 64"),
        ("Random Rotation", "随机旋转 10 度以内"),
        ("Color Jitter", "亮度、对比度、饱和度扰动"),
        ("Random Affine", "小幅平移和缩放"),
        ("Normalize", "按固定均值和标准差归一化"),
    ]
    add_table(doc, ["处理步骤", "说明"], rows, [2600, 6760])
    add_caption(doc, "表 2  数据预处理与增强策略")


def add_baseline(doc):
    doc.add_heading("3. 基础交通标志识别模型", level=1)
    doc.add_heading("3.1 模型结构", level=2)
    add_body_paragraph(
        doc,
        "本项目选用 ResNet18 作为基础分类模型。ResNet 的核心思想是引入残差连接，使网络能够学习输入与输出之间的残差映射，从而缓解深层网络训练中的梯度消失和退化问题。由于 GTSRB 图像尺寸较小，本项目将 ResNet18 第一层卷积调整为 3 x 3 卷积，并去除初始最大池化层，以保留更多空间细节；最后将全连接层输出类别数改为 43。",
    )
    rows = [
        ("模型", "ResNet18"),
        ("输入尺寸", "64 x 64"),
        ("类别数", "43"),
        ("优化器", "AdamW"),
        ("学习率", "0.001"),
        ("权重衰减", "0.0001"),
        ("学习率调度", "Cosine"),
        ("Epoch", "30"),
        ("Batch Size", "256"),
        ("损失函数", "CrossEntropyLoss"),
    ]
    add_table(doc, ["参数", "设置"], rows, [2600, 6760])
    add_caption(doc, "表 3  基础识别模型训练配置")

    doc.add_heading("3.2 基础识别结果", level=2)
    metrics = json.load(open(TABLE_DIR / "table_03_baseline_test_metrics.json", encoding="utf-8"))
    rows = [
        ("Test Accuracy", f"{metrics['accuracy']:.4f}"),
        ("Precision Macro", f"{metrics['precision_macro']:.4f}"),
        ("Recall Macro", f"{metrics['recall_macro']:.4f}"),
        ("F1 Macro", f"{metrics['f1_macro']:.4f}"),
        ("参数量", metrics["parameter_count"]),
        ("最优 Epoch", metrics["best_epoch"]),
    ]
    add_table(doc, ["指标", "数值"], rows, [3000, 6360])
    add_caption(doc, "表 4  ResNet18 基础识别测试结果")
    add_figure(doc, "fig_05_baseline_loss_curve.png", "图 3  ResNet18 训练损失曲线", width=5.65)
    add_figure(doc, "fig_06_baseline_accuracy_curve.png", "图 4  ResNet18 训练与验证准确率曲线", width=5.65)
    add_figure(doc, "fig_07_baseline_confusion_matrix.png", "图 5  ResNet18 在测试集上的混淆矩阵", width=5.85)
    add_body_paragraph(
        doc,
        "从实验结果可以看出，ResNet18 在正常测试集上取得 97.81% 的准确率，说明基础识别模型能够较好完成交通标志分类任务。错误样例显示，模型在低光照、模糊、遮挡或类别相似情况下仍可能发生误判。",
    )
    add_figure(doc, "fig_10_baseline_wrong_samples.png", "图 6  基础识别模型错误分类样例", width=6.25)


def add_attack(doc):
    doc.add_heading("4. 对抗攻击方法与实验结果", level=1)
    doc.add_heading("4.1 对抗样本定义", level=2)
    add_body_paragraph(
        doc,
        "对抗样本是在原始样本上加入微小扰动后得到的输入。设原始图像为 x，真实标签为 y，分类模型为 f，攻击目标是在扰动大小受限的情况下构造 x_adv，使模型预测错误。epsilon 用于控制扰动强度，epsilon 越大，攻击能力通常越强，但图像变化也可能越明显。",
    )
    doc.add_heading("4.2 FGSM 与 PGD 攻击", level=2)
    add_body_paragraph(
        doc,
        "FGSM 是一种单步梯度符号攻击方法，沿损失函数对输入图像梯度的符号方向添加扰动，使模型损失快速增大。PGD 是一种迭代式攻击方法，每一步沿梯度方向更新对抗样本，并将扰动投影回 epsilon 范围内。相比 FGSM，PGD 通常攻击能力更强。",
    )
    rows = [
        ("目标模型", "ResNet18 best_model.pth"),
        ("测试样本", "从测试集分层抽样 3000 张"),
        ("FGSM epsilon", "0.005, 0.01, 0.02, 0.03, 0.05"),
        ("PGD epsilon", "0.005, 0.01, 0.02, 0.03, 0.05"),
        ("PGD alpha", "0.005"),
        ("PGD steps", "7"),
    ]
    add_table(doc, ["实验项", "设置"], rows, [2600, 6760])
    add_caption(doc, "表 5  对抗攻击实验设置")

    attack_df = pd.read_csv(TABLE_DIR / "table_06_attack_metrics.csv")
    fgsm_rows = []
    pgd_rows = []
    for _, row in attack_df.iterrows():
        values = [
            f"{row['epsilon']:.3f}",
            f"{row['clean_accuracy']:.4f}",
            f"{row['adversarial_accuracy']:.4f}",
            f"{row['attack_success_rate']:.4f}",
        ]
        if row["attack"] == "fgsm":
            fgsm_rows.append(values)
        else:
            pgd_rows.append(values)
    add_table(doc, ["Epsilon", "Clean Acc.", "Adv. Acc.", "Attack Success"], fgsm_rows, [1600, 2500, 2500, 2760])
    add_caption(doc, "表 6  FGSM 攻击实验结果")
    add_table(doc, ["Epsilon", "Clean Acc.", "Adv. Acc.", "Attack Success"], pgd_rows, [1600, 2500, 2500, 2760])
    add_caption(doc, "表 7  PGD 攻击实验结果")
    add_figure(doc, "fig_11_attack_accuracy_curve.png", "图 7  不同 epsilon 下的对抗样本准确率变化", width=5.85)
    add_figure(doc, "fig_12_attack_success_curve.png", "图 8  不同 epsilon 下的攻击成功率变化", width=5.85)
    add_figure(doc, "fig_14_fgsm_eps003_triplets.png", "图 9  FGSM epsilon=0.03 对抗样本展示", width=6.25)
    add_figure(doc, "fig_15_pgd_eps003_triplets.png", "图 10  PGD epsilon=0.03 对抗样本展示", width=6.25)
    add_body_paragraph(
        doc,
        "实验结果表明，随着 epsilon 增大，模型在对抗样本上的准确率明显下降。FGSM 在 epsilon=0.03 时将模型准确率从 97.93% 降至 67.10%，PGD 在相同 epsilon 下进一步降至 60.00%。这说明即使模型在正常测试集上具有较高准确率，也可能受到较小输入扰动的显著影响。PGD 通过多次迭代寻找更有效的扰动方向，因此在多数扰动强度下比 FGSM 更强。",
    )


def add_defense(doc):
    doc.add_heading("5. 输入预处理防御方法与实验结果", level=1)
    doc.add_heading("5.1 防御方法", level=2)
    add_body_paragraph(
        doc,
        "输入预处理防御的基本思想是在图像送入模型前进行一定变换，以削弱对抗扰动的影响。本项目选取三种轻量方法：Gaussian Blur 用于平滑局部噪声，Median Filter 用于抑制局部异常像素，JPEG Compression 通过压缩去除部分高频扰动。这些方法不需要重新训练模型，部署简单，适合作为轻量防御基线。",
    )
    rows = [
        ("Gaussian Blur", "使用高斯核平滑图像，削弱局部噪声"),
        ("Median Filter", "使用中值滤波抑制异常像素"),
        ("JPEG Compression", "通过有损压缩去除部分高频扰动"),
    ]
    add_table(doc, ["防御方法", "作用机制"], rows, [2600, 6760])
    add_caption(doc, "表 8  输入预处理防御方法")

    doc.add_heading("5.2 防御实验结果", level=2)
    defense_df = pd.read_csv(TABLE_DIR / "table_07_input_defense_metrics.csv")
    for eps in [0.03, 0.05]:
        rows = []
        for attack in ["fgsm", "pgd"]:
            sub = defense_df[(defense_df["attack"] == attack) & (defense_df["epsilon"].round(2) == eps)]
            before = sub.iloc[0]["adversarial_accuracy_before_defense"]
            values = {r["defense"]: r["defended_accuracy"] for _, r in sub.iterrows()}
            rows.append(
                [
                    attack.upper(),
                    f"{before:.4f}",
                    f"{values['gaussian_blur']:.4f}",
                    f"{values['median_filter']:.4f}",
                    f"{values['jpeg_compression']:.4f}",
                ]
            )
        add_table(
            doc,
            ["Attack", "Before", "Gaussian", "Median", "JPEG"],
            rows,
            [1500, 1800, 2000, 1900, 2160],
        )
        add_caption(doc, f"表 {9 if eps == 0.03 else 10}  epsilon={eps:.2f} 下的输入预处理防御结果")

    add_figure(doc, "fig_16_fgsm_input_defense_curve.png", "图 11  FGSM 攻击下输入预处理防御效果", width=5.85)
    add_figure(doc, "fig_17_pgd_input_defense_curve.png", "图 12  PGD 攻击下输入预处理防御效果", width=5.85)
    add_figure(doc, "fig_18_input_defense_accuracy_bar.png", "图 13  epsilon=0.03 下不同输入防御方法对比", width=6.25)
    add_figure(doc, "fig_19_input_defense_examples.png", "图 14  输入预处理防御恢复案例", width=6.25)
    add_body_paragraph(
        doc,
        "三种输入预处理方法均能在一定程度上恢复模型准确率，其中 JPEG Compression 效果最明显。在 PGD epsilon=0.03 场景下，模型准确率从 60.30% 提升至 79.43%；在 PGD epsilon=0.05 场景下，从 54.47% 提升至 71.73%。JPEG Compression 效果较好的原因可能是对抗扰动常包含较多高频细节，而 JPEG 压缩会丢弃部分高频信息，从而削弱扰动影响。",
    )
    add_body_paragraph(
        doc,
        "不过，输入预处理防御并不能完全恢复到正常测试集准确率。这说明简单输入变换只能作为基础防御方法，面对更强攻击或自适应攻击时仍存在局限。",
    )


def add_implementation_and_process(doc):
    doc.add_heading("6. 程序实现与使用说明", level=1)
    doc.add_heading("6.1 项目结构", level=2)
    rows = [
        ("configs/", "实验配置文件"),
        ("src/", "源代码"),
        ("results/", "完整实验结果"),
        ("reports/", "报告素材、结果汇总和 Word 报告"),
        ("docs/", "实验计划和规范文档"),
    ]
    add_table(doc, ["目录", "说明"], rows, [2400, 6960])
    add_caption(doc, "表 11  项目目录结构")

    doc.add_heading("6.2 核心代码文件", level=2)
    rows = [
        ("src/data/check_gtsrb.py", "数据集检查、统计表和样例图生成"),
        ("src/models/classifiers.py", "SimpleCNN 和 ResNet18 模型定义"),
        ("src/train_classifier.py", "基础分类模型训练与评估"),
        ("src/attacks/methods.py", "FGSM 和 PGD 攻击方法"),
        ("src/evaluate_attacks.py", "对抗攻击评估与可视化"),
        ("src/evaluate_input_defense.py", "输入预处理防御评估"),
    ]
    add_table(doc, ["文件", "功能"], rows, [3300, 6060])
    add_caption(doc, "表 12  核心代码文件说明")

    doc.add_heading("6.3 运行命令", level=2)
    for cmd in [
        "python -m src.data.check_gtsrb --config configs/dataset_check.yaml",
        "python -m src.train_classifier --config configs/baseline_resnet18.yaml",
        "python -m src.evaluate_attacks --config configs/attack_fgsm_pgd.yaml",
        "python -m src.evaluate_input_defense --config configs/defense_input_preprocessing.yaml",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(cmd)
        set_run_font(r, name="Consolas", east_asia="Microsoft YaHei", size=9.5, color=RGBColor(40, 40, 40))

    doc.add_heading("7. 实验过程记录与问题分析", level=1)
    doc.add_heading("7.1 已完成实验", level=2)
    rows = [
        ("E00", "GTSRB 数据集检查与可视化", "已完成"),
        ("E01", "ResNet18 基础交通标志识别", "已完成"),
        ("E02", "FGSM / PGD 对抗攻击", "已完成"),
        ("E03", "输入预处理防御", "已完成"),
        ("E04", "对抗训练防御", "代码已准备，未纳入正式结果"),
    ]
    add_table(doc, ["编号", "实验内容", "状态"], rows, [1200, 5200, 2960])
    add_caption(doc, "表 13  实验完成情况")

    doc.add_heading("7.2 遇到的问题", level=2)
    for item in [
        "PGD 攻击计算耗时较长。PGD 每个 batch 需要多次前向和反向传播，在全测试集、多 epsilon 设置下耗时较长。为保证实验可控，本项目采用分层抽样的 3000 张测试图进行攻击与防御评估，并保存抽样索引，确保不同实验在同一批样本上比较。",
        "可视化图片排版问题。初始样例宫格图中类别名称较长，出现文字拥挤。后续调整了单元格宽度和标签显示方式，使样例图更适合报告展示。",
        "输入防御存在性能上限。输入预处理可以提升对抗样本下的准确率，但也可能损失原图中的边缘和纹理信息，无法完全恢复到 clean accuracy。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("8. 创新点、不足与改进方向", level=1)
    doc.add_heading("8.1 创新点", level=2)
    for item in [
        "没有停留在普通交通标志分类，而是进一步分析模型安全性和鲁棒性。",
        "构建了“基础识别 - 对抗攻击 - 防御评估”的完整实验闭环。",
        "同时比较 FGSM 和 PGD 两种攻击方法，分析不同 epsilon 下的性能变化。",
        "保存原图、扰动图、对抗样本和防御恢复案例，增强实验展示性。",
        "使用同一批分层抽样测试样本比较攻击和防御结果，提高实验公平性。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("8.2 不足与改进方向", level=2)
    for item in [
        "输入预处理防御属于轻量防御，面对自适应攻击时可能效果下降。",
        "对抗训练防御代码已创建，但尚未完整跑完并验证，因此未纳入正式结果。",
        "尚未加入 Grad-CAM 可解释性分析，后续可从注意力区域角度解释攻击效果。",
        "尚未完成交互式 Demo，后续可使用 Streamlit 实现上传图片、生成攻击和防御结果的可视化系统。",
    ]:
        add_bullet(doc, item)


def add_summary_and_roles(doc):
    doc.add_heading("9. 小组分工", level=1)
    rows = [
        ("成员 A", "数据集处理、数据可视化、结果整理", "20%"),
        ("成员 B", "ResNet18 模型训练与基础识别实验", "20%"),
        ("成员 C", "FGSM / PGD 攻击实现与实验", "20%"),
        ("成员 D", "输入预处理防御实现与对比分析", "20%"),
        ("成员 E", "实验报告、PPT、答辩材料整理", "20%"),
    ]
    add_table(doc, ["成员", "工作内容", "贡献占比"], rows, [1600, 6000, 1760])
    add_caption(doc, "表 14  小组分工模板，提交前请按实际成员修改")

    doc.add_heading("10. 总结", level=1)
    add_body_paragraph(
        doc,
        "本项目完成了基于 ResNet18 的交通标志识别模型，并在 GTSRB 测试集上取得 97.81% 的准确率。在此基础上，项目实现了 FGSM 和 PGD 两种对抗攻击方法，实验发现模型在对抗扰动下准确率显著下降，说明深度学习交通标志识别模型存在明显鲁棒性问题。进一步地，项目使用 Gaussian Blur、Median Filter 和 JPEG Compression 三种输入预处理方法进行防御实验，结果表明 JPEG Compression 可以较明显提升对抗样本下的识别准确率。",
    )
    add_body_paragraph(
        doc,
        "整体来看，本项目从普通图像分类任务扩展到模型安全性分析，形成了识别、攻击、防御和可视化的完整实验链路，具有较好的技术深度和展示效果。后续可以继续补充对抗训练、Grad-CAM 可解释性分析和交互式 Demo，以进一步提高项目完整度。",
    )


def build_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_abstract(doc)
    add_background(doc)
    add_dataset(doc)
    add_baseline(doc)
    add_attack(doc)
    add_defense(doc)
    add_implementation_and_process(doc)
    add_summary_and_roles(doc)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_docx()
    print(path)
